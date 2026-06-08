from __future__ import annotations

import csv
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parent
REPO = ROOT.parent
RESULTS = REPO / "results"
FIGURES = ROOT / "figures_v2"

OUT_MD = ROOT / "Sealing_Jutsu_Research_Paper_v3.md"
OUT_DOCX = ROOT / "Sealing_Jutsu_Research_Paper_v3.docx"
OUT_PDF = ROOT / "Sealing_Jutsu_Research_Paper_v3.pdf"

TITLE = "Intent-Bound Memory Authorization Against Persistent Agent Poisoning"
AUTHORS = [
    ("Akshay Jain", "Independent Researcher", "jakshay623@gmail.com"),
    ("Sujit Suryawanshi", "Independent Researcher", "sujitsuryawanshi987@gmail.com"),
]

BLUE = "1F4D78"
DARK = "111827"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GRID = "D0D5DD"


def pct(value: float | str, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except (TypeError, ValueError):
        return str(value)


def rows(path: str) -> list[dict[str, str]]:
    with (RESULTS / path).open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def row_by(path: str, key: str) -> dict[str, dict[str, str]]:
    return {row[key]: row for row in rows(path)}


def metric(row: dict[str, str], name: str, default: str = "0") -> str:
    return row.get(name, default)


def build_context() -> dict[str, object]:
    workflow = row_by("workflow_corpus_test_split_summary.csv", "agent")
    lifecycle = row_by("memory_lifecycle_gap_summary.csv", "agent")
    agentdojo = row_by("converted_agentdojo_all_summary.csv", "agent")
    injec_all = row_by("converted_injecagent_all_summary.csv", "agent")
    injec_dh = row_by("converted_injecagent_dh_summary.csv", "agent")
    injec_ds = row_by("converted_injecagent_ds_summary.csv", "agent")
    live_llm = row_by("gap_fix_medium_live_llm_summary.csv", "condition")
    live_models = rows("gap_fix_medium_live_llm_model_summary.csv")
    calibration = rows("current_main_threshold_calibration.csv")
    generated = row_by("all_scenarios_generated_holdout_summary.csv", "agent")
    advanced = row_by("advanced_attack_suite_summary.csv", "agent")
    return {
        "workflow": workflow,
        "lifecycle": lifecycle,
        "agentdojo": agentdojo,
        "injec_all": injec_all,
        "injec_dh": injec_dh,
        "injec_ds": injec_ds,
        "live_llm": live_llm,
        "live_models": live_models,
        "calibration": calibration,
        "generated": generated,
        "advanced": advanced,
    }


def md_table(headers: list[str], data: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |"]
    out.append("| " + " | ".join(["---"] * len(headers)) + " |")
    out.extend("| " + " | ".join(row) + " |" for row in data)
    return "\n".join(out)


def paper_blocks(ctx: dict[str, object]) -> list[tuple[str, object]]:
    workflow: dict[str, dict[str, str]] = ctx["workflow"]  # type: ignore[assignment]
    lifecycle: dict[str, dict[str, str]] = ctx["lifecycle"]  # type: ignore[assignment]
    agentdojo: dict[str, dict[str, str]] = ctx["agentdojo"]  # type: ignore[assignment]
    injec_all: dict[str, dict[str, str]] = ctx["injec_all"]  # type: ignore[assignment]
    injec_dh: dict[str, dict[str, str]] = ctx["injec_dh"]  # type: ignore[assignment]
    injec_ds: dict[str, dict[str, str]] = ctx["injec_ds"]  # type: ignore[assignment]
    live_llm: dict[str, dict[str, str]] = ctx["live_llm"]  # type: ignore[assignment]
    live_models: list[dict[str, str]] = ctx["live_models"]  # type: ignore[assignment]
    generated: dict[str, dict[str, str]] = ctx["generated"]  # type: ignore[assignment]
    advanced: dict[str, dict[str, str]] = ctx["advanced"]  # type: ignore[assignment]

    def summary_row(source: dict[str, dict[str, str]], agent: str) -> list[str]:
        r = source[agent]
        return [
            agent,
            pct(metric(r, "attack_success_rate_mean")),
            pct(metric(r, "unauthorized_risky_action_rate_mean")),
            pct(metric(r, "benign_accuracy_mean")),
            pct(metric(r, "false_positive_rate_mean")),
            pct(metric(r, "poison_sealing_rate_mean")),
        ]

    def influence_row(corpus: str, source: dict[str, dict[str, str]]) -> list[str]:
        out_mod = source["output_moderation"]
        cap = source["intent_capsules"]
        return [
            corpus,
            pct(metric(out_mod, "attack_success_rate_mean")),
            pct(metric(out_mod, "poison_influence_rate_mean")),
            pct(metric(cap, "attack_success_rate_mean")),
            pct(metric(cap, "poison_influence_rate_mean")),
            pct(metric(cap, "poison_sealing_rate_mean")),
        ]

    model_rows = []
    for item in live_models:
        if item["condition"] == "capsule_filtered_prompt":
            model_rows.append(
                [
                    item["model"],
                    str(int(float(item["cases"]))),
                    pct(item["planner_attack_success_rate"]),
                    pct(item["attack_success_rate"]),
                    pct(item["unauthorized_risky_action_rate"]),
                    pct(item["raw_parse_error_rate"]),
                ]
            )

    return [
        ("abstract", (
            "Long-term-memory LLM agents can remember user preferences, tool outputs, web observations, OCR text, summaries, and prior "
            "experiences across sessions. That persistence improves personalization, but it also creates a durable poisoning channel: "
            "a memory inserted during one interaction can later be retrieved as ordinary context and steer recommendations, plans, or "
            "tool calls. This paper argues that persistent agent poisoning is not only a malicious-content detection problem; it is an "
            "authority-control problem. We introduce intent-bound memory authorization, implemented in the open-source CapsuleGuard "
            "prototype in the Sealing-Jutsu repository. CapsuleGuard compiles stored memories into capsules carrying source type, topic "
            "scope, denied actions, authority score, influence budget, verification count, lineage, freshness, and status. Retrieval is "
            "treated as relevance, not authorization: before a memory can shape planning or action, its capsule contract must permit "
            "that influence for the current intent and risk level. Across the held-out workflow corpus, intent capsules reduced attack "
            f"success from {pct(metric(workflow['ambient_memory'], 'attack_success_rate_mean'))} for ambient memory to "
            f"{pct(metric(workflow['intent_capsules'], 'attack_success_rate_mean'))}, while preserving "
            f"{pct(metric(workflow['intent_capsules'], 'benign_accuracy_mean'))} benign accuracy and "
            f"{pct(metric(workflow['intent_capsules'], 'false_positive_rate_mean'))} false positives. The newest evaluation adds "
            "poison influence rate, showing that output moderation can block final actions while still allowing poisoned memory to steer "
            "planning. On converted AgentDojo and InjecAgent traces, output moderation and semantic judges reach 0.00% final ASR but "
            "retain 30.00%-90.62% poison influence, whereas intent capsules reach 0.00% ASR and 0.00% poison influence. A medium live "
            "LLM planner run across llama3, mistral, and phi3 shows 22.22% ambient final attack success and 0.00% capsule-filtered final "
            "attack success. These results support least-privilege memory authorization as a promising defense layer for persistent agent "
            "memory poisoning, while the paper limits its claim to the tested threat model and prototype evaluation."
        )),
        ("keywords", "LLM agents, memory poisoning, prompt injection, agent security, least privilege, authorization, long-term memory, tool safety"),
        ("h1", "1. Introduction"),
        ("p", (
            "LLM agents increasingly use persistent memory to store user preferences, previous task outcomes, tool results, retrieved web "
            "content, OCR observations, and generated summaries. This gives agents continuity across sessions, but it also creates a "
            "cross-session attack surface. If an attacker can cause a misleading or malicious memory to be stored today, that memory can "
            "remain dormant and later reappear as apparently normal context."
        )),
        ("p", (
            "The central weakness is ambient memory authority. Many systems ask whether a memory is relevant enough to retrieve, but they "
            "do not separately ask what that memory is authorized to influence. Once placed into a planner context, a poisoned memory may "
            "act like a fact, preference, instruction, tool hint, or prior experience. Persistent poisoning succeeds because relevance and "
            "authority are collapsed into one step."
        )),
        ("p", (
            "This paper presents intent-bound memory authorization. The design turns each memory into a bounded security object. A capsule "
            "may still be retrieved for context, but it cannot influence a recommendation, plan, or tool call unless its contract allows "
            "that influence. The contribution is not another keyword filter or output judge; it is a least-privilege control plane for "
            "memory influence."
        )),
        ("h2", "Contributions"),
        ("bullets", [
            "A reformulation of persistent agent poisoning as an ambient-authority failure in memory systems.",
            "A capsule schema binding memory to source, topic, denied action, authority, influence budget, verification, lineage, freshness, and status.",
            "A runtime authorization gate that separates retrieval from influence and blocks unauthorized memory before planning or action execution.",
            "A reproducible Python prototype with baseline agents, stress suites, trace-corpus loading, signed provenance, vector-style retrieval, safe tool traces, and live LLM planner support.",
            "A new poison influence rate metric that distinguishes final output blocking from planner compromise.",
            "Evaluation across synthetic workflow tasks, lifecycle-gap attacks, converted AgentDojo/InjecAgent traces, ablations, calibration sweeps, and local live LLM planners.",
        ]),
        ("h1", "2. Background and Research Gap"),
        ("p", (
            "Prior work has shown that agent memory can be poisoned through direct memory injection, multimodal content, corrupted web "
            "state, poisoned tool outputs, query-only interaction, poisoned knowledge bases, and experience retrieval. The common lesson "
            "is that stored content can become a long-lived steering mechanism. Prompt sanitization and one-time filtering are not enough "
            "when the attack can sleep in memory and trigger later."
        )),
        ("table", (
            "Table 1. Defensive gap addressed by intent-bound capsules.",
            ["Defense", "Question answered", "What remains open"],
            [
                ["Keyword filtering", "Does this text look suspicious?", "Benign-phrased poison survives."],
                ["Prompt sanitization", "Can injection text be removed at intake?", "Dormant memory can bypass later checks."],
                ["Provenance only", "Where did the memory come from?", "A legitimate-looking source is not the same as action authority."],
                ["Trust-score retrieval", "Which memory should rank higher?", "High trust does not define allowed influence."],
                ["Output moderation", "Is the final action unsafe?", "Poisoned memory may already have shaped planning."],
                ["Intent capsules", "Is this memory authorized for this intent and action?", "A complementary authorization layer, not a universal jailbreak detector."],
            ],
        )),
        ("h1", "3. Threat Model"),
        ("figure", (FIGURES / "figure2_threat_model_v2.png", "Figure 1. Threat model and trust boundaries.")),
        ("p", (
            "The adversary is a memory-poisoning attacker, not a code-execution attacker. The attacker can cause untrusted or weakly "
            "trusted content to enter the agent memory lifecycle, wait for that content to be retrieved in a later task, and attempt to "
            "turn retrieved data into planning authority."
        )),
        ("table", (
            "Table 2. Formal attacker capabilities and limits.",
            ["Capability", "In scope", "Explicit limit"],
            [
                ["Injection channel", "Web pages, tool outputs, OCR text, alt text, user-visible documents, experience logs, summaries, or memory imports.", "No direct edit of system prompts, policy code, capsule gate, or benchmark labels."],
                ["Persistence", "Delayed-trigger or cross-session poison stored now and activated later.", "Cannot delete audit records or force status transitions to be ignored."],
                ["Semantic adaptation", "Paraphrases, split payloads, benign wording, retrieval-collision text.", "Cannot forge cryptographic provenance, verified writer identity, or independent corroboration."],
                ["Tool/OCR manipulation", "Malicious instructions inside tool-return text or OCR-visible content.", "Cannot compromise the tool runtime, OCR binary, or external credentials."],
                ["Goal", "Alter recommendations, planning path, memory-derived preference, or medium/high-risk tool action.", "Not focused on denial-of-service, model-weight compromise, or arbitrary host code execution."],
            ],
        )),
        ("p", (
            "Protected assets are recommendation integrity, user preference integrity, action safety, memory lifecycle integrity, and "
            "auditability. The trusted computing base for the prototype includes the capsule compiler, policy gate, and authorization "
            "rules. The memory database, retriever, OCR pipeline, browser/tool outputs, and agent summaries are treated as partially "
            "untrusted. The core security boundary is between retrieved memory as relevance evidence and retrieved memory as authority."
        )),
        ("h1", "4. Design: Intent-Bound Memory Capsules"),
        ("figure", (FIGURES / "figure1_architecture_v2.png", "Figure 2. CapsuleGuard architecture.")),
        ("p", (
            "A capsule is a memory record plus an influence contract. The contract does not merely describe provenance; it defines what "
            "the memory is permitted to affect. A web page may be useful context without authorizing a purchase. A tool result may describe "
            "state without rewriting user preference. A generated summary may aid recall while inheriting the authority limits of its sources."
        )),
        ("table", (
            "Table 3. Capsule fields and security role.",
            ["Field", "Security role"],
            [
                ["source_type", "Classifies origin such as user declaration, verified record, web content, tool output, OCR, summary, or experience."],
                ["kind", "Separates facts, preferences, observations, experiences, tool results, and directive-like records."],
                ["allowed_topics", "Limits influence to matching user intents."],
                ["denied_actions", "Blocks actions the memory can never authorize."],
                ["source_authority", "Caps how much trust a source class can carry."],
                ["influence_budget", "Prevents one memory from dominating a plan by relevance alone."],
                ["verification_count", "Records corroboration from independent evidence."],
                ["lineage", "Prevents summaries and derived memories from laundering weak authority."],
                ["freshness/status", "Reduces stale influence and keeps sealed/rejected states auditable."],
            ],
        )),
        ("figure", (FIGURES / "figure3_trust_rules_v2.png", "Figure 3. Memory trust and influence rule set.")),
        ("h2", "Authorization Rules"),
        ("bullets", [
            "Topic scope: a memory may influence only intents with sufficient topic overlap.",
            "Action denial: weak external sources cannot authorize dangerous actions by default.",
            "Authority floors: medium/high-risk actions require higher source authority.",
            "Directive sealing: instruction-like memories are kept for audit but removed from planning eligibility.",
            "Influence budget: each memory has bounded decision weight.",
            "Evidence quorum: risky actions require independent support across source and writer diversity.",
            "Lineage and freshness: derived or stale records cannot silently become stronger than their sources.",
        ]),
        ("h1", "5. Prototype Implementation"),
        ("p", (
            "CapsuleGuard is implemented as a Python research prototype in the Sealing-Jutsu repository. It includes capsule compilation, "
            "policy gates, baseline agents, attack scenario generation, converted trace-corpus loading, signed provenance, vector-style "
            "retrieval, safe tool simulation, JSONL traces, CSV reports, charts, and a live LLM planner harness. The full test suite now "
            "contains 163 passing tests."
        )),
        ("table", (
            "Table 4. Prototype components.",
            ["Component", "Purpose"],
            [
                ["capsule_guard/compiler.py", "Builds capsules from raw memory seeds."],
                ["capsule_guard/policy.py", "Enforces topic, action, quorum, and authorization gates."],
                ["capsule_guard/evaluation.py", "Computes ASR, risky action, benign accuracy, false positives, sealing, and poison influence."],
                ["capsule_guard/workflow_corpus.py", "Loads workflow and external trace corpora."],
                ["capsule_guard/llm_planner.py", "Runs strict planner-schema LLM experiments."],
                ["run_capsuleguard.py", "Main benchmark runner and report generator."],
            ],
        )),
        ("h1", "6. Evaluation Methodology"),
        ("p", (
            "The evaluation compares intent capsules against ambient memory, keyword filtering, quarantine-only retrieval, trust-score "
            "retrieval, provenance-only retrieval, counterfactual memory, output moderation, and semantic output judging. Metrics include "
            "attack success rate (ASR), unauthorized risky action rate, poison influence rate, benign accuracy, false positive rate, sealing "
            "rate, and latency."
        )),
        ("p", (
            "Poison influence rate is the critical new metric. It counts poisoned cases where poisoned memory is used by the planner and "
            "the planner selects the attacker target before any final output gate. This separates defenses that block the final action from "
            "defenses that prevent poisoned memory from shaping the plan."
        )),
        ("h1", "7. Results"),
        ("h2", "7.1 Held-Out Workflow Corpus"),
        ("figure", (FIGURES / "figure4_workflow_results_v2.png", "Figure 4. Held-out workflow-corpus results.")),
        ("table", (
            "Table 5. Held-out workflow-corpus test split.",
            ["Agent", "ASR", "Risky action", "Benign accuracy", "FPR", "Sealing"],
            [
                summary_row(workflow, "ambient_memory"),
                summary_row(workflow, "keyword_filter"),
                summary_row(workflow, "output_moderation"),
                summary_row(workflow, "provenance_only"),
                summary_row(workflow, "counterfactual_memory"),
                summary_row(workflow, "intent_capsules"),
            ],
        )),
        ("p", (
            "Intent capsules are the only tested defense in this split that reaches 0.00% ASR, 0.00% risky action, 100.00% benign "
            "accuracy, and 0.00% false positives at the same time. Counterfactual memory reduces ASR but sacrifices utility, while "
            "provenance and trust-score retrieval help but do not close the residual risk."
        )),
        ("h2", "7.2 Memory Lifecycle Gap"),
        ("table", (
            "Table 6. Poison influence exposes late-blocking defenses.",
            ["Agent", "ASR", "Risky action", "Poison influence", "Benign accuracy"],
            [
                [
                    "output_moderation",
                    pct(metric(lifecycle["output_moderation"], "attack_success_rate_mean")),
                    pct(metric(lifecycle["output_moderation"], "unauthorized_risky_action_rate_mean")),
                    pct(metric(lifecycle["output_moderation"], "poison_influence_rate_mean")),
                    pct(metric(lifecycle["output_moderation"], "benign_accuracy_mean")),
                ],
                [
                    "semantic_output_judge",
                    pct(metric(lifecycle["semantic_output_judge"], "attack_success_rate_mean")),
                    pct(metric(lifecycle["semantic_output_judge"], "unauthorized_risky_action_rate_mean")),
                    pct(metric(lifecycle["semantic_output_judge"], "poison_influence_rate_mean")),
                    pct(metric(lifecycle["semantic_output_judge"], "benign_accuracy_mean")),
                ],
                [
                    "ablation_no_denied_actions",
                    pct(metric(lifecycle["ablation_no_denied_actions"], "attack_success_rate_mean")),
                    pct(metric(lifecycle["ablation_no_denied_actions"], "unauthorized_risky_action_rate_mean")),
                    pct(metric(lifecycle["ablation_no_denied_actions"], "poison_influence_rate_mean")),
                    pct(metric(lifecycle["ablation_no_denied_actions"], "benign_accuracy_mean")),
                ],
                [
                    "intent_capsules",
                    pct(metric(lifecycle["intent_capsules"], "attack_success_rate_mean")),
                    pct(metric(lifecycle["intent_capsules"], "unauthorized_risky_action_rate_mean")),
                    pct(metric(lifecycle["intent_capsules"], "poison_influence_rate_mean")),
                    pct(metric(lifecycle["intent_capsules"], "benign_accuracy_mean")),
                ],
            ],
        )),
        ("p", (
            "This benchmark answers the reviewer question: if output moderation can reach low ASR, why use memory authorization? The "
            "answer is that output moderation can leave poisoned influence intact. Intent capsules are designed to block unauthorized "
            "memory influence before the planner or final action gate can be steered."
        )),
        ("h2", "7.3 Converted AgentDojo/InjecAgent Trace Corpora"),
        ("table", (
            "Table 7. External converted trace-corpus readout.",
            ["Corpus", "Output-mod ASR", "Output-mod influence", "Capsule ASR", "Capsule influence", "Capsule sealing"],
            [
                influence_row("AgentDojo all", agentdojo),
                influence_row("InjecAgent all", injec_all),
                influence_row("InjecAgent DH", injec_dh),
                influence_row("InjecAgent DS", injec_ds),
            ],
        )),
        ("p", (
            "On the converted corpora, output moderation and semantic judges often block the final direct action, producing 0.00% ASR. "
            "However, poison influence remains high because the poisoned memory still reaches the planner and selects the attacker target. "
            "Intent capsules reach 0.00% ASR and 0.00% poison influence across all converted splits, while sealing a large fraction of "
            "poisoned records before planning."
        )),
        ("h2", "7.4 Stress Suites and Ablations"),
        ("figure", (FIGURES / "figure5_stress_asr_v2.png", "Figure 5. Stress-suite attack success rates.")),
        ("table", (
            "Table 8. Stress-suite examples.",
            ["Suite", "Ambient ASR", "Provenance ASR", "Capsule ASR", "Capsule benign"],
            [
                [
                    "Generated holdout",
                    pct(metric(generated["ambient_memory"], "attack_success_rate_mean")),
                    pct(metric(generated["provenance_only"], "attack_success_rate_mean")),
                    pct(metric(generated["intent_capsules"], "attack_success_rate_mean")),
                    pct(metric(generated["intent_capsules"], "benign_accuracy_mean")),
                ],
                [
                    "Advanced suite",
                    pct(metric(advanced["ambient_memory"], "attack_success_rate_mean")),
                    pct(metric(advanced["provenance_only"], "attack_success_rate_mean")),
                    pct(metric(advanced["intent_capsules"], "attack_success_rate_mean")),
                    pct(metric(advanced["intent_capsules"], "benign_accuracy_mean")),
                ],
            ],
        )),
        ("p", (
            "Ablations show that low ASR alone can be misleading. Removing topic scope can reduce attacks by blocking too much benign "
            "utility. Removing denied-action controls can hide final ASR while allowing high poison influence. The complete design matters "
            "because it combines influence authorization, action denial, evidence quorum, and utility preservation."
        )),
        ("h2", "7.5 Live LLM Planner Check"),
        ("figure", (FIGURES / "figure7_live_llm_planner_v2.png", "Figure 6. Live LLM planner check.")),
        ("table", (
            "Table 9. Medium live LLM workflow-corpus run.",
            ["Condition", "Rows", "Planner tempted", "Final ASR", "Risky action", "Raw parse error"],
            [
                [
                    "ambient_prompt",
                    str(int(float(live_llm["ambient_prompt"]["cases"]))),
                    pct(live_llm["ambient_prompt"]["planner_attack_success_rate"]),
                    pct(live_llm["ambient_prompt"]["attack_success_rate"]),
                    pct(live_llm["ambient_prompt"]["unauthorized_risky_action_rate"]),
                    pct(live_llm["ambient_prompt"]["raw_parse_error_rate"]),
                ],
                [
                    "capsule_filtered_prompt",
                    str(int(float(live_llm["capsule_filtered_prompt"]["cases"]))),
                    pct(live_llm["capsule_filtered_prompt"]["planner_attack_success_rate"]),
                    pct(live_llm["capsule_filtered_prompt"]["attack_success_rate"]),
                    pct(live_llm["capsule_filtered_prompt"]["unauthorized_risky_action_rate"]),
                    pct(live_llm["capsule_filtered_prompt"]["raw_parse_error_rate"]),
                ],
            ],
        )),
        ("table", (
            "Table 10. Defended live LLM result by model.",
            ["Model", "Rows", "Planner tempted", "Final ASR", "Risky action", "Raw parse error"],
            model_rows,
        )),
        ("p", (
            "The live LLM experiment is not the largest benchmark; it is a realism check. Across llama3, mistral, and phi3, the ambient "
            "prompt produces 22.22% final attack success. Capsule-filtered planning still shows 2.78% planner temptation, which is useful "
            "evidence that the planner can be attracted to poison, but final authorization reduces accepted ASR and risky action to 0.00%."
        )),
        ("h2", "7.6 Threshold Calibration"),
        ("p", (
            "A 16-point current-main sweep varied medium-risk quorum and topic-scope thresholds. In the simulator, all swept settings "
            "reported 0.00% ASR, 0.00% risky action, 100.00% benign accuracy, and 0.00% false positives. This closes the prototype-level "
            "hand-tuned-only gap, but it does not replace calibration on larger external benign and adversarial workloads."
        )),
        ("h1", "8. Discussion"),
        ("p", (
            "The strongest result is not simply that intent capsules reach 0.00% ASR in the tested settings. The stronger result is that "
            "the method targets a different security property than output moderation. Output moderation can block a visible dangerous "
            "action after the planner has already been steered. Intent-bound capsules deny unauthorized memory influence before the "
            "planning authority is granted."
        )),
        ("p", (
            "The converted-corpus results make this distinction concrete: output moderation and semantic output judging show 0.00% ASR "
            "while retaining 30.00%-90.62% poison influence, depending on corpus. That means the final action was blocked, but memory "
            "compromise still occurred. For memory-security research, this matters because persistent influence can accumulate, shape "
            "intermediate decisions, or reappear in later workflows even when one final action gate succeeds."
        )),
        ("h1", "9. Limitations and Threats to Validity"),
        ("bullets", [
            "The high-volume planner is still mostly deterministic, although live LLM planner evidence is now included.",
            "Several corpora are synthetic, generated, or converted rather than collected from long-running deployed users.",
            "The converted corpora are useful for external-style stress, but conversion choices may simplify real workflow complexity.",
            "Raw multimodal hidden-pixel/OCR pipelines are not fully evaluated; OCR-style text is covered more strongly than image forensics.",
            "The prototype assumes policy code, cryptographic attestations, and verified writer identity are not compromised.",
            "The result is not a proof that all memory poisoning is solved; it is evidence for least-privilege memory authorization under the stated threat model.",
            "Frontier paid API models and larger autonomous red-team loops remain future validation work.",
        ]),
        ("h1", "10. Conclusion"),
        ("p", (
            "Persistent agent poisoning succeeds when stored memory is treated as ordinary context with ambient authority. Intent-bound "
            "memory authorization removes that ambient authority by requiring each memory to prove what it is allowed to influence. In "
            "the current CapsuleGuard prototype, this approach reaches 0.00% ASR and 0.00% poison influence across the tested converted "
            "trace corpora, preserves benign utility in the held-out workflow corpus, and blocks final attack success in live local LLM "
            "planner checks. The results do not establish universal security, but they do support the core research claim: memory retrieval "
            "should be separated from memory authority, and agent memory should be governed by least-privilege influence contracts."
        )),
        ("h1", "References"),
        ("references", [
            "M. A. Ferrag, A. Lakas, N. Tihanyi, and M. Debbah, Securing LLM agents: From prompt sanitization to autonomous red teaming and beyond, Internet of Things and Cyber-Physical Systems, 2025.",
            "J. Qian, Visual Inception: Compromising Long-term Planning in Agentic Recommenders via Multimodal Memory Poisoning, arXiv, 2026.",
            "X. Yang, Y. He, S. Ji, B. Hooi, and J. S. Dong, Zombie Agents: Persistent Control of Self-Evolving LLM Agents via Self-Reinforcing Injections, arXiv, 2026.",
            "Z. Chen, Z. Xiang, C. Xiao, D. Song, and B. Li, AgentPoison: Red-teaming LLM Agents via Poisoning Memory or Knowledge Bases, arXiv, 2024.",
            "A. S. Patlan, A. Hebbar, P. Viswanath, and P. Mittal, Context manipulation attacks: Web agents are susceptible to corrupted memory, arXiv, 2025.",
            "Z. Xu, X. Zhu, Y. Yao, M. Xue, and Y. Song, From Storage to Steering: Memory Control Flow Attacks on LLM Agents, arXiv, 2026.",
            "H. Tian, Z. Sha, J. Wang, Y. Liu, Z. Huang, and X. Huang, InjecMEM: Memory Injection Attack on LLM Agent Memory Systems, OpenReview, 2025.",
            "V. Torra and M. Bras-Amoros, Memory poisoning and secure multi-agent systems, arXiv, 2026.",
            "B. D. Sunil et al., Memory Poisoning Attack and Defense on Memory Based LLM-Agents, arXiv, 2026.",
            "S. S. Srivastava and H. He, MemoryGraft: Persistent Compromise of LLM Agents via Poisoned Experience Retrieval, arXiv, 2025.",
            "S. Dong et al., Memory Injection Attacks on LLM Agents via Query-Only Interaction, arXiv, 2025.",
            "OWASP Foundation, OWASP Top 10 for Large Language Model Applications, 2025.",
            "A. Jain and S. Suryawanshi, Sealing-Jutsu: CapsuleGuard prototype, https://github.com/Mr-Akuma/Sealing-Jutsu.",
        ]),
    ]


def write_markdown(blocks: list[tuple[str, object]]) -> None:
    parts = [
        f"# {TITLE}",
        "",
        "Akshay Jain, Independent Researcher, jakshay623@gmail.com",
        "",
        "Sujit Suryawanshi, Independent Researcher, sujitsuryawanshi987@gmail.com",
        "",
        "Draft generated from Sealing-Jutsu repository artifacts, June 2026.",
        "",
    ]
    for kind, payload in blocks:
        if kind == "abstract":
            parts += ["## Abstract", str(payload), ""]
        elif kind == "keywords":
            parts += [f"**Index Terms:** {payload}", ""]
        elif kind == "h1":
            parts += [f"## {payload}", ""]
        elif kind == "h2":
            parts += [f"### {payload}", ""]
        elif kind == "p":
            parts += [str(payload), ""]
        elif kind == "bullets":
            parts += [f"- {item}" for item in payload] + [""]
        elif kind == "table":
            caption, headers, data = payload
            parts += [f"**{caption}**", "", md_table(headers, data), ""]
        elif kind == "figure":
            path, caption = payload
            rel = path.relative_to(ROOT).as_posix()
            parts += [f"![{caption}]({rel})", "", f"*{caption}*", ""]
        elif kind == "references":
            for idx, ref in enumerate(payload, 1):
                parts.append(f"{idx}. {ref}")
            parts.append("")
    OUT_MD.write_text("\n".join(parts), encoding="utf-8")


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str = DARK) -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, caption: str, headers: list[str], data: list[list[str]]) -> None:
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    set_run_font(run, 9.5, True, BLUE)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        shade_cell(hdr[idx], LIGHT_BLUE)
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        r = hdr[idx].paragraphs[0].add_run(header)
        set_run_font(r, 8.5, True, DARK)
    for row in data:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            r = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(r, 8.3, False, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, 10.2, False, DARK)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, 14 if level == 1 else 12, True, BLUE)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(5.9))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    run = cp.add_run(caption)
    set_run_font(run, 8.8, False, MUTED)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run, 10.0, False, DARK)


def build_docx(blocks: list[tuple[str, object]]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(TITLE)
    set_run_font(tr, 18, True, BLUE)

    for name, affiliation, email in AUTHORS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{name}\n{affiliation}\n{email}")
        set_run_font(r, 9.5, False, DARK)
        p.paragraph_format.space_after = Pt(3)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run("Draft generated from Sealing-Jutsu repository artifacts, June 2026")
    set_run_font(mr, 9, False, MUTED)

    for kind, payload in blocks:
        if kind == "abstract":
            add_heading(doc, "Abstract", 1)
            add_para(doc, str(payload))
        elif kind == "keywords":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r1 = p.add_run("Index Terms: ")
            set_run_font(r1, 10.2, True, DARK)
            r2 = p.add_run(str(payload))
            set_run_font(r2, 10.2, False, DARK)
        elif kind == "h1":
            add_heading(doc, str(payload), 1)
        elif kind == "h2":
            add_heading(doc, str(payload), 2)
        elif kind == "p":
            add_para(doc, str(payload))
        elif kind == "bullets":
            add_bullets(doc, payload)
        elif kind == "table":
            caption, headers, data = payload
            add_table(doc, caption, headers, data)
        elif kind == "figure":
            path, caption = payload
            add_figure(doc, path, caption)
        elif kind == "references":
            for idx, ref in enumerate(payload, 1):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(-0.22)
                p.paragraph_format.left_indent = Inches(0.22)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"[{idx}] {ref}")
                set_run_font(run, 8.8, False, DARK)

    doc.save(OUT_DOCX)


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "PaperTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=17,
            leading=20,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceAfter=10,
        ),
        "author": ParagraphStyle(
            "AuthorBlock",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{DARK}"),
            spaceAfter=5,
        ),
        "meta": ParagraphStyle(
            "Meta",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=10,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=12,
        ),
        "h1": ParagraphStyle(
            "Heading1Custom",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=12,
            spaceAfter=5,
        ),
        "h2": ParagraphStyle(
            "Heading2Custom",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "BodyCustom",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.6,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "SmallCustom",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=7.2,
            leading=8.5,
            spaceAfter=2,
        ),
        "caption": ParagraphStyle(
            "CaptionCustom",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8,
            leading=9.5,
            alignment=TA_CENTER,
            textColor=colors.HexColor(f"#{MUTED}"),
            spaceAfter=6,
        ),
        "table_caption": ParagraphStyle(
            "TableCaptionCustom",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=8.3,
            leading=10,
            textColor=colors.HexColor(f"#{BLUE}"),
            spaceBefore=5,
            spaceAfter=3,
        ),
    }


def pdf_para(text: str, style: ParagraphStyle) -> Paragraph:
    safe = (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    return Paragraph(safe, style)


def pdf_markup(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(str(text), style)


def pdf_table(caption: str, headers: list[str], data: list[list[str]], styles: dict[str, ParagraphStyle], width: float) -> list[object]:
    flow: list[object] = [pdf_para(caption, styles["table_caption"])]
    rows_pdf = [[pdf_para(header, styles["small"]) for header in headers]]
    rows_pdf.extend([[pdf_para(cell, styles["small"]) for cell in row] for row in data])
    col_width = width / max(len(headers), 1)
    table = Table(rows_pdf, colWidths=[col_width] * len(headers), repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_BLUE}")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor(f"#{DARK}")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{GRID}")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{LIGHT_GRAY}")]),
            ]
        )
    )
    flow += [table, Spacer(1, 6)]
    return flow


def pdf_bullets(items: Iterable[str], styles: dict[str, ParagraphStyle]) -> ListFlowable:
    return ListFlowable(
        [ListItem(pdf_para(item, styles["body"]), leftIndent=12) for item in items],
        bulletType="bullet",
        leftIndent=18,
        bulletFontName="Helvetica",
        bulletFontSize=7,
    )


def pdf_figure(path: Path, caption: str, styles: dict[str, ParagraphStyle], width: float) -> list[object]:
    if not path.exists():
        return []
    img = PdfImage(str(path))
    scale = min(width / img.drawWidth, 3.75 * inch / img.drawHeight)
    img.drawWidth *= scale
    img.drawHeight *= scale
    img.hAlign = "CENTER"
    return [Spacer(1, 5), img, pdf_para(caption, styles["caption"])]


def add_page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
    canvas.drawRightString(doc.pagesize[0] - doc.rightMargin, 0.38 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(blocks: list[tuple[str, object]]) -> None:
    doc = SimpleDocTemplate(
        str(OUT_PDF),
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.6 * inch,
        title=TITLE,
        author=", ".join(name for name, _, _ in AUTHORS),
    )
    styles = pdf_styles()
    width = doc.width
    story: list[object] = [pdf_para(TITLE, styles["title"])]
    for name, affiliation, email in AUTHORS:
        story.append(pdf_markup(f"{name}<br/>{affiliation}<br/>{email}", styles["author"]))
    story.append(pdf_para("Draft generated from Sealing-Jutsu repository artifacts, June 2026", styles["meta"]))

    for kind, payload in blocks:
        if kind == "abstract":
            story.append(pdf_para("Abstract", styles["h1"]))
            story.append(pdf_para(str(payload), styles["body"]))
        elif kind == "keywords":
            story.append(pdf_markup(f"<b>Index Terms:</b> {payload}", styles["body"]))
        elif kind == "h1":
            story.append(pdf_para(str(payload), styles["h1"]))
        elif kind == "h2":
            story.append(pdf_para(str(payload), styles["h2"]))
        elif kind == "p":
            story.append(pdf_para(str(payload), styles["body"]))
        elif kind == "bullets":
            story.append(pdf_bullets(payload, styles))
            story.append(Spacer(1, 4))
        elif kind == "table":
            caption, headers, data = payload
            story.extend(pdf_table(caption, headers, data, styles, width))
        elif kind == "figure":
            path, caption = payload
            story.extend(pdf_figure(path, caption, styles, width))
        elif kind == "references":
            for idx, ref in enumerate(payload, 1):
                story.append(pdf_para(f"[{idx}] {ref}", styles["small"]))

    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def main() -> None:
    ctx = build_context()
    blocks = paper_blocks(ctx)
    write_markdown(blocks)
    build_docx(blocks)
    build_pdf(blocks)
    print(f"wrote {OUT_MD}")
    print(f"wrote {OUT_DOCX}")
    print(f"wrote {OUT_PDF}")


if __name__ == "__main__":
    main()
