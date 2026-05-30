from __future__ import annotations

import csv
import math
import textwrap
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
REPO_ROOT_CANDIDATES = [
    ROOT.parent,
    Path(r"C:\Users\User\Music\Agent-Poisoning-Research-FINAL"),
]
REPO_ROOT = next((path for path in REPO_ROOT_CANDIDATES if (path / "results").exists()), REPO_ROOT_CANDIDATES[-1])
SOURCE_RESULTS = REPO_ROOT / "results"
OUT_MD = ROOT / "Intent_Bound_Memory_Capsules_Submission_Draft_v2.md"
OUT_DOCX = ROOT / "Intent_Bound_Memory_Capsules_Submission_Draft_v2.docx"
OUT_TABLES = ROOT / "paper_result_tables_v2.csv"
FIG_DIR = ROOT / "figures_v2"

TITLE = "Intent-Bound Memory Capsules: Least-Privilege Authorization for Persistent LLM Agent Memory"
AUTHOR = "Akshay Jain"
AFFILIATION = "Independent Researcher, India"
DATE_LINE = "Draft built from local artifacts, May 2026"

NAVY = "#0B2545"
BLUE = "#1F4D78"
LIGHT_BLUE = "#E8EEF5"
GREEN = "#176B3A"
RED = "#B42318"
ORANGE = "#B86B00"
GRAY = "#667085"
LIGHT_GRAY = "#F2F4F7"
GRID = "#D0D5DD"
BLACK = "#111827"


def hex_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def short_pct(value: float) -> str:
    return f"{value * 100:.0f}"


def read_summary(name: str) -> dict[str, dict[str, float]]:
    path = SOURCE_RESULTS / name
    rows: dict[str, dict[str, float]] = {}
    with path.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            agent = row["agent"]
            rows[agent] = {}
            for key, value in row.items():
                if key == "agent":
                    continue
                try:
                    rows[agent][key] = float(value)
                except (TypeError, ValueError):
                    pass
    return rows


def read_metric_rows(name: str, key_field: str) -> dict[str, dict[str, float | str]]:
    path = SOURCE_RESULTS / name
    rows: dict[str, dict[str, float | str]] = {}
    with path.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            key = row[key_field]
            parsed: dict[str, float | str] = {}
            for field, value in row.items():
                if field == key_field:
                    continue
                try:
                    parsed[field] = float(value)
                except (TypeError, ValueError):
                    parsed[field] = value
            rows[key] = parsed
    return rows


def read_metric_list(name: str) -> list[dict[str, float | str]]:
    path = SOURCE_RESULTS / name
    rows: list[dict[str, float | str]] = []
    with path.open(newline="", encoding="utf-8") as handle:
        reader = csv.DictReader(handle)
        for row in reader:
            parsed: dict[str, float | str] = {}
            for field, value in row.items():
                try:
                    parsed[field] = float(value)
                except (TypeError, ValueError):
                    parsed[field] = value
            rows.append(parsed)
    return rows


def load_all_results() -> dict[str, Any]:
    workflow = read_summary("workflow_corpus_test_split_summary.csv")
    scenarios = {
        "Generated holdout": read_summary("all_scenarios_generated_holdout_summary.csv"),
        "Attacker-generated": read_summary("all_scenarios_attacker_generated_summary.csv"),
        "Multimodal/OCR": read_summary("all_scenarios_multimodal_summary.csv"),
        "Trusted-source compromise": read_summary("all_scenarios_trusted_source_compromise_summary.csv"),
        "Advanced suite": read_summary("advanced_attack_suite_summary.csv"),
    }
    llm_summary = read_metric_rows("gap_fix_medium_live_llm_summary.csv", "condition")
    llm_model_summary = read_metric_list("gap_fix_medium_live_llm_model_summary.csv")
    llm_gap_report = read_metric_list("gap_fix_medium_live_llm_gap_report.csv")[0]
    high_cost_summary = read_metric_rows("high_cost_local_smoke_summary.csv", "condition")
    high_cost_statistics = read_metric_list("high_cost_local_smoke_statistics.csv")[0]
    return {
        "workflow": workflow,
        "scenarios": scenarios,
        "llm_summary": llm_summary,
        "llm_model_summary": llm_model_summary,
        "llm_gap_report": llm_gap_report,
        "high_cost_summary": high_cost_summary,
        "high_cost_statistics": high_cost_statistics,
    }


def write_tables_csv(results: dict[str, Any]) -> None:
    rows: list[dict[str, str]] = []
    workflow = results["workflow"]
    for agent, metrics in workflow.items():
        rows.append(
            {
                "table": "held_out_workflow",
                "scenario": "workflow_corpus_test_split",
                "agent": agent,
                "attack_success_rate": str(metrics.get("attack_success_rate_mean", "")),
                "unauthorized_risky_action_rate": str(metrics.get("unauthorized_risky_action_rate_mean", "")),
                "benign_accuracy": str(metrics.get("benign_accuracy_mean", "")),
                "poison_sealing_rate": str(metrics.get("poison_sealing_rate_mean", "")),
                "false_positive_rate": str(metrics.get("false_positive_rate_mean", "")),
            }
        )
    for scenario, data in results["scenarios"].items():
        for agent in ("ambient_memory", "provenance_only", "intent_capsules"):
            metrics = data[agent]
            rows.append(
                {
                    "table": "stress_scenarios",
                    "scenario": scenario,
                    "agent": agent,
                    "attack_success_rate": str(metrics.get("attack_success_rate_mean", "")),
                    "unauthorized_risky_action_rate": str(metrics.get("unauthorized_risky_action_rate_mean", "")),
                    "benign_accuracy": str(metrics.get("benign_accuracy_mean", "")),
                    "poison_sealing_rate": str(metrics.get("poison_sealing_rate_mean", "")),
                    "false_positive_rate": str(metrics.get("false_positive_rate_mean", "")),
                }
            )
    for condition, metrics in results["llm_summary"].items():
        rows.append(
            {
                "table": "medium_live_llm",
                "scenario": "workflow_corpus_live_planner",
                "agent": condition,
                "cases": str(metrics.get("cases", "")),
                "planner_attack_success_rate": str(metrics.get("planner_attack_success_rate", "")),
                "attack_success_rate": str(metrics.get("attack_success_rate", "")),
                "unauthorized_risky_action_rate": str(metrics.get("unauthorized_risky_action_rate", "")),
                "raw_parse_error_rate": str(metrics.get("raw_parse_error_rate", "")),
                "parse_error_rate": str(metrics.get("parse_error_rate", "")),
            }
        )
    for condition, metrics in results["high_cost_summary"].items():
        rows.append(
            {
                "table": "high_cost_local_smoke",
                "scenario": "high_cost_local_smoke",
                "agent": condition,
                "cases": str(metrics.get("cases", "")),
                "planner_attack_success_rate": str(metrics.get("planner_attack_success_rate", "")),
                "attack_success_rate": str(metrics.get("attack_success_rate", "")),
                "unauthorized_risky_action_rate": str(metrics.get("unauthorized_risky_action_rate", "")),
                "raw_parse_error_rate": str(metrics.get("raw_parse_error_rate", "")),
                "parse_error_rate": str(metrics.get("parse_error_rate", "")),
            }
        )
    with OUT_TABLES.open("w", newline="", encoding="utf-8") as handle:
        fieldnames = sorted({field for row in rows for field in row.keys()})
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    width: int,
    fnt: ImageFont.ImageFont,
    fill: str = BLACK,
    line_gap: int = 5,
) -> int:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if text_size(draw, trial, fnt)[0] <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    line_h = text_size(draw, "Ag", fnt)[1] + line_gap
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=hex_rgb(fill))
        y += line_h
    return y


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = BLUE, width: int = 3) -> None:
    draw.line([start, end], fill=hex_rgb(color), width=width)
    sx, sy = start
    ex, ey = end
    ang = math.atan2(ey - sy, ex - sx)
    size = 10
    pts = [
        (ex, ey),
        (ex - size * math.cos(ang - 0.45), ey - size * math.sin(ang - 0.45)),
        (ex - size * math.cos(ang + 0.45), ey - size * math.sin(ang + 0.45)),
    ]
    draw.polygon(pts, fill=hex_rgb(color))


def box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    subtitle: str = "",
    fill: str = LIGHT_BLUE,
    outline: str = BLUE,
    title_fill: str = NAVY,
) -> None:
    draw.rounded_rectangle(xy, radius=10, fill=hex_rgb(fill), outline=hex_rgb(outline), width=2)
    x1, y1, x2, y2 = xy
    title_font = font(15, True)
    sub_font = font(12)
    tw, th = text_size(draw, title, title_font)
    draw.text((x1 + (x2 - x1 - tw) / 2, y1 + 16), title, font=title_font, fill=hex_rgb(title_fill))
    if subtitle:
        wrapped_y = y1 + 42
        draw_wrapped(draw, (x1 + 14, wrapped_y), subtitle, x2 - x1 - 28, sub_font, GRAY, 3)


def new_canvas(w: int = 1500, h: int = 850) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    return img, draw


def draw_title(draw: ImageDraw.ImageDraw, title: str, subtitle: str | None = None) -> None:
    draw.text((55, 34), title, font=font(24, True), fill=hex_rgb(NAVY))
    if subtitle:
        draw.text((55, 68), subtitle, font=font(14), fill=hex_rgb(GRAY))
    draw.line((55, 98, 1445, 98), fill=hex_rgb(GRID), width=2)


def figure_architecture() -> Path:
    img, draw = new_canvas()
    draw_title(draw, "Figure 1. Intent-Bound Memory Capsule Architecture", "Retrieval finds candidate memories; authorization decides what they may influence.")
    y_top = 160
    nodes = [
        ((70, y_top, 250, y_top + 90), "Raw memory", "user, web, OCR, tool, summary"),
        ((315, y_top, 515, y_top + 90), "Capsule compiler", "classify, scope, cap, seal"),
        ((580, y_top, 790, y_top + 90), "Provenance ledger", "source, writer, lineage, time"),
        ((855, y_top, 1045, y_top + 90), "Capsule store", "active plus sealed records"),
    ]
    for xy, t, s in nodes:
        box(draw, xy, t, s)
    for a, b in zip(nodes, nodes[1:]):
        arrow(draw, (a[0][2], y_top + 45), (b[0][0], y_top + 45))

    lower = [
        ((170, 430, 370, 520), "Retrieval", "semantic relevance only"),
        ((445, 430, 675, 520), "Authorization gates", "topic, action, authority, quorum"),
        ((755, 430, 970, 520), "Planner", "only authorized context"),
        ((1045, 430, 1245, 520), "Tool/action gate", "risk-aware execution"),
    ]
    for xy, t, s in lower:
        box(draw, xy, t, s, fill="#F5F8FF")
    arrow(draw, (950, y_top + 90), (270, 430), BLUE)
    for a, b in zip(lower, lower[1:]):
        arrow(draw, (a[0][2], 475), (b[0][0], 475))

    box(draw, (445, 620, 675, 710), "Sealed / blocked", "stored for audit, not planning", fill="#FEF3F2", outline=RED, title_fill=RED)
    draw.line((560, 520, 560, 620), fill=hex_rgb(RED), width=3)
    arrow(draw, (560, 620), (560, 620), RED)

    chips = ["status", "topic scope", "source authority", "denied actions", "influence budget", "verification", "lineage", "freshness"]
    x, y = 430, 555
    for chip in chips:
        w = text_size(draw, chip, font(12))[0] + 34
        draw.rounded_rectangle((x, y, x + w, y + 30), radius=15, fill=hex_rgb("#FFFFFF"), outline=hex_rgb(GRID))
        draw.text((x + 17, y + 8), chip, font=font(12), fill=hex_rgb(BLUE))
        x += w + 12
        if x > 1300:
            x = 430
            y += 42
    out = FIG_DIR / "figure1_architecture_v2.png"
    img.save(out, quality=95)
    return out


def figure_threat_model() -> Path:
    img, draw = new_canvas()
    draw_title(draw, "Figure 2. Threat Model and Trust Boundaries", "Persistent poisoning enters as data but can become unauthorized influence if retrieval is treated as authority.")
    sources = [
        ("Web content", "indirect prompt/data poison"),
        ("Tool output", "result manipulation"),
        ("OCR/document text", "multimodal hidden text"),
        ("Agent summary", "self-generated memory"),
        ("Experience log", "retrieved prior episode"),
        ("Query-only input", "gradual memory shaping"),
    ]
    y = 150
    for label, sub in sources:
        box(draw, (70, y, 300, y + 66), label, sub, fill="#FFF7ED", outline=ORANGE, title_fill="#7C2D12")
        arrow(draw, (300, y + 33), (510, 340), ORANGE, 2)
        y += 86
    box(draw, (510, 275, 750, 405), "Long-term memory", "persistent store across sessions", fill="#F8FAFC", outline=GRAY)
    box(draw, (840, 255, 1090, 425), "Capsule authorization boundary", "least-privilege influence checks", fill=LIGHT_BLUE, outline=BLUE)
    box(draw, (1190, 275, 1410, 405), "Planner and tools", "recommend, email, purchase, call APIs", fill="#ECFDF3", outline=GREEN, title_fill=GREEN)
    arrow(draw, (750, 340), (840, 340), BLUE)
    arrow(draw, (1090, 340), (1190, 340), GREEN)
    box(draw, (840, 515, 1090, 635), "Blocked influence", "relevance without authority is not injected", fill="#FEF3F2", outline=RED, title_fill=RED)
    draw.line((965, 425, 965, 515), fill=hex_rgb(RED), width=3)
    arrow(draw, (965, 515), (965, 515), RED)
    trust = [
        "Policy config is trusted",
        "Raw external content is untrusted",
        "Memory database is persistent",
        "Tool side effects are high impact",
    ]
    y = 660
    draw.text((70, y), "Trust-boundary assumptions", font=font(18, True), fill=hex_rgb(NAVY))
    y += 35
    for item in trust:
        draw.ellipse((80, y + 5, 91, y + 16), fill=hex_rgb(BLUE))
        draw.text((105, y), item, font=font(15), fill=hex_rgb(BLACK))
        y += 30
    out = FIG_DIR / "figure2_threat_model_v2.png"
    img.save(out, quality=95)
    return out


def grouped_bar_chart(
    title: str,
    subtitle: str,
    labels: list[str],
    series: list[tuple[str, str, list[float]]],
    out: Path,
    h: int = 850,
) -> Path:
    img, draw = new_canvas(1500, h)
    draw_title(draw, title, subtitle)
    left, right, top, bottom = 115, 1420, 165, h - 150
    draw.line((left, top, left, bottom), fill=hex_rgb(GRAY), width=2)
    draw.line((left, bottom, right, bottom), fill=hex_rgb(GRAY), width=2)
    for i in range(0, 101, 25):
        y = bottom - (bottom - top) * i / 100
        draw.line((left, y, right, y), fill=hex_rgb("#EAECF0"), width=1)
        draw.text((55, y - 8), f"{i}%", font=font(12), fill=hex_rgb(GRAY))
    group_w = (right - left) / len(labels)
    bar_w = min(34, group_w / (len(series) + 2))
    for i, label in enumerate(labels):
        center = left + group_w * (i + 0.5)
        start = center - (len(series) * bar_w + (len(series) - 1) * 8) / 2
        for j, (sname, color, values) in enumerate(series):
            value = max(0.0, min(values[i], 1.0))
            x1 = start + j * (bar_w + 8)
            x2 = x1 + bar_w
            y1 = bottom - (bottom - top) * value
            draw.rectangle((x1, y1, x2, bottom), fill=hex_rgb(color))
            draw.text((x1 - 3, y1 - 24), short_pct(value), font=font(11, True), fill=hex_rgb(color))
        lines = textwrap.wrap(label, width=17)
        y_label = bottom + 24
        for line in lines[:3]:
            tw, _ = text_size(draw, line, font(11))
            draw.text((center - tw / 2, y_label), line, font=font(11), fill=hex_rgb(BLACK))
            y_label += 16
    lx = left
    ly = 118
    for name, color, _ in series:
        draw.rectangle((lx, ly, lx + 18, ly + 18), fill=hex_rgb(color))
        draw.text((lx + 26, ly + 1), name, font=font(13), fill=hex_rgb(BLACK))
        lx += text_size(draw, name, font(13))[0] + 80
    img.save(out, quality=95)
    return out


def build_figures(results: dict[str, Any]) -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    paths["architecture"] = figure_architecture()
    paths["threat"] = figure_threat_model()

    workflow = results["workflow"]
    labels = ["Ambient", "Keyword", "Provenance", "Counterfactual", "Capsules"]
    agents = ["ambient_memory", "keyword_filter", "provenance_only", "counterfactual_memory", "intent_capsules"]
    paths["rules"] = figure_rules()

    paths["workflow"] = grouped_bar_chart(
        "Figure 4. Held-Out Workflow Corpus Results",
        "Capsules block tested poisoning while preserving benign accuracy; counterfactual defense reduces ASR but harms utility.",
        labels,
        [
            ("ASR", RED, [workflow[a]["attack_success_rate_mean"] for a in agents]),
            ("Risky action", ORANGE, [workflow[a]["unauthorized_risky_action_rate_mean"] for a in agents]),
            ("Benign", GREEN, [workflow[a]["benign_accuracy_mean"] for a in agents]),
        ],
        FIG_DIR / "figure4_workflow_results_v2.png",
    )

    scenario_labels = list(results["scenarios"].keys())
    scenario_data = results["scenarios"]
    paths["stress"] = grouped_bar_chart(
        "Figure 5. Attack Success Rate Across Stress Suites",
        "The capsule defense reports zero attack success in these simulator suites; baselines remain exposed.",
        scenario_labels,
        [
            ("Ambient", RED, [scenario_data[s]["ambient_memory"]["attack_success_rate_mean"] for s in scenario_labels]),
            ("Provenance", "#0E7490", [scenario_data[s]["provenance_only"]["attack_success_rate_mean"] for s in scenario_labels]),
            ("Capsules", GREEN, [scenario_data[s]["intent_capsules"]["attack_success_rate_mean"] for s in scenario_labels]),
        ],
        FIG_DIR / "figure5_stress_asr_v2.png",
    )

    ablation_agents = ["intent_capsules", "ablation_no_topic_scope", "ablation_no_denied_actions", "ablation_no_quorum"]
    ablation_labels = ["Full capsules", "No topic scope", "No denied actions", "No quorum"]
    paths["ablation"] = grouped_bar_chart(
        "Figure 6. Ablation Results on Held-Out Workflow Corpus",
        "Removing topic scope keeps ASR low in this split but damages utility, so ASR alone is not enough.",
        ablation_labels,
        [
            ("ASR", RED, [workflow[a]["attack_success_rate_mean"] for a in ablation_agents]),
            ("Benign", GREEN, [workflow[a]["benign_accuracy_mean"] for a in ablation_agents]),
            ("FPR", ORANGE, [workflow[a]["false_positive_rate_mean"] for a in ablation_agents]),
        ],
        FIG_DIR / "figure6_ablation_v2.png",
    )

    llm = results["llm_summary"]
    llm_labels = ["Ambient prompt", "Capsule-filtered prompt"]
    llm_keys = ["ambient_prompt", "capsule_filtered_prompt"]
    paths["live_llm"] = grouped_bar_chart(
        "Figure 7. Live LLM Planner Check",
        "The LLM can be tempted before authorization, but capsule-filtered final ASR and risky action remain zero in the medium run.",
        llm_labels,
        [
            ("Planner tempted", ORANGE, [llm[k]["planner_attack_success_rate"] for k in llm_keys]),
            ("Final ASR", RED, [llm[k]["attack_success_rate"] for k in llm_keys]),
            ("Risky action", "#7C2D12", [llm[k]["unauthorized_risky_action_rate"] for k in llm_keys]),
        ],
        FIG_DIR / "figure7_live_llm_planner_v2.png",
    )

    make_contact_sheet(paths)
    return paths


def figure_rules() -> Path:
    img, draw = new_canvas(1500, 900)
    draw_title(draw, "Figure 3. Memory Trust and Influence Rule Set", "Memories are allowed to inform only the planning/action classes supported by their capsule contract.")
    headers = ["Source / memory class", "Default authority", "Allowed influence", "Main restrictions"]
    rows = [
        ("Verified record", "0.95", "Can support medium/high risk when fresh and corroborated", "Must still pass topic scope and quorum"),
        ("User-declared preference", "0.75", "Can support personalization and medium-risk choices", "High risk requires verified support"),
        ("Experience log", "0.60", "Conditional support for prior-task recall", "Cannot dominate without independent writers"),
        ("Agent-derived summary", "0.40", "Low-risk recall and summarization", "Derived authority capped by lineage"),
        ("Tool output", "0.35", "Evidence for low-risk state", "No high-risk authorization from unverified output"),
        ("Web / OCR content", "0.20-0.25", "Context only", "Denied from dangerous actions; must be corroborated"),
    ]
    x0, y0 = 55, 145
    widths = [310, 180, 440, 430]
    row_h = 84
    x = x0
    for idx, header in enumerate(headers):
        draw.rectangle((x, y0, x + widths[idx], y0 + 52), fill=hex_rgb(NAVY), outline=hex_rgb(NAVY))
        draw.text((x + 14, y0 + 16), header, font=font(14, True), fill=(255, 255, 255))
        x += widths[idx]
    y = y0 + 52
    for r, row in enumerate(rows):
        fill = "#FFFFFF" if r % 2 == 0 else "#F8FAFC"
        x = x0
        for idx, value in enumerate(row):
            draw.rectangle((x, y, x + widths[idx], y + row_h), fill=hex_rgb(fill), outline=hex_rgb(GRID))
            draw_wrapped(draw, (x + 14, y + 14), value, widths[idx] - 28, font(14 if idx else 15, idx == 0), BLACK)
            x += widths[idx]
        y += row_h
    draw.text((55, 800), "Risk floors: LOW 0.00, MEDIUM 0.55, HIGH 0.85. Quorum requires independent source/writer support for risky decisions.", font=font(15, True), fill=hex_rgb(BLUE))
    out = FIG_DIR / "figure3_trust_rules_v2.png"
    img.save(out, quality=95)
    return out


def make_contact_sheet(paths: dict[str, Path]) -> Path:
    thumbs: list[tuple[str, Image.Image]] = []
    for label, path in paths.items():
        img = Image.open(path).convert("RGB")
        img.thumbnail((610, 360))
        thumbs.append((path.name, img.copy()))
    rows = math.ceil(len(thumbs) / 2)
    w, h = 1300, 110 + rows * 385
    sheet = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(sheet)
    draw.text((35, 24), "Submission Draft Figure QA Contact Sheet", font=font(24), fill=hex_rgb(BLACK))
    x_positions = [35, 680]
    y = 82
    for idx, (name, img) in enumerate(thumbs):
        col = idx % 2
        if col == 0 and idx > 0:
            y += 385
        x = x_positions[col]
        draw.text((x, y), name, font=font(17), fill=hex_rgb(BLACK))
        sheet.paste(img, (x, y + 34))
    out = FIG_DIR / "figure_contact_sheet_v2.png"
    sheet.save(out, quality=95)
    return out


def workflow_table(results: dict[str, Any]) -> list[list[str]]:
    workflow = results["workflow"]
    order = [
        ("ambient_memory", "Ambient memory"),
        ("keyword_filter", "Keyword filter"),
        ("provenance_only", "Provenance only"),
        ("counterfactual_memory", "Counterfactual memory"),
        ("intent_capsules", "Intent-bound capsules"),
    ]
    rows = []
    for key, label in order:
        row = workflow[key]
        rows.append(
            [
                label,
                pct(row["attack_success_rate_mean"]),
                pct(row["unauthorized_risky_action_rate_mean"]),
                pct(row["benign_accuracy_mean"]),
                pct(row["false_positive_rate_mean"]),
                pct(row["poison_sealing_rate_mean"]),
            ]
        )
    return rows


def stress_table(results: dict[str, Any]) -> list[list[str]]:
    rows = []
    for scenario, data in results["scenarios"].items():
        rows.append(
            [
                scenario,
                pct(data["ambient_memory"]["attack_success_rate_mean"]),
                pct(data["provenance_only"]["attack_success_rate_mean"]),
                pct(data["intent_capsules"]["attack_success_rate_mean"]),
                pct(data["intent_capsules"]["poison_sealing_rate_mean"]),
            ]
        )
    return rows


def ablation_table(results: dict[str, Any]) -> list[list[str]]:
    workflow = results["workflow"]
    order = [
        ("intent_capsules", "Full capsule defense"),
        ("ablation_no_topic_scope", "No topic scope"),
        ("ablation_no_denied_actions", "No denied actions"),
        ("ablation_no_quorum", "No quorum"),
    ]
    rows = []
    for key, label in order:
        row = workflow[key]
        rows.append(
            [
                label,
                pct(row["attack_success_rate_mean"]),
                pct(row["benign_accuracy_mean"]),
                pct(row["false_positive_rate_mean"]),
                pct(row["poison_sealing_rate_mean"]),
                capsule_interpretation(key),
            ]
        )
    return rows


def live_llm_table(results: dict[str, Any]) -> list[list[str]]:
    llm = results["llm_summary"]
    order = [
        ("ambient_prompt", "Ambient prompt"),
        ("capsule_filtered_prompt", "Capsule-filtered prompt"),
    ]
    rows = []
    for key, label in order:
        row = llm[key]
        rows.append(
            [
                label,
                str(int(row["cases"])),
                pct(row["planner_attack_success_rate"]),
                pct(row["attack_success_rate"]),
                pct(row["unauthorized_risky_action_rate"]),
                pct(row["raw_parse_error_rate"]),
                pct(row["parse_error_rate"]),
            ]
        )
    return rows


def live_llm_model_table(results: dict[str, Any]) -> list[list[str]]:
    rows = []
    for row in results["llm_model_summary"]:
        if row["condition"] != "capsule_filtered_prompt":
            continue
        rows.append(
            [
                str(row["model"]),
                str(int(row["cases"])),
                pct(row["planner_attack_success_rate"]),
                pct(row["attack_success_rate"]),
                pct(row["unauthorized_risky_action_rate"]),
                pct(row["raw_parse_error_rate"]),
            ]
        )
    return rows


def high_cost_smoke_table(results: dict[str, Any]) -> list[list[str]]:
    summary = results["high_cost_summary"]
    stats = results["high_cost_statistics"]
    return [
        [
            "Ambient prompt",
            str(int(summary["ambient_prompt"]["cases"])),
            pct(summary["ambient_prompt"]["attack_success_rate"]),
            pct(summary["ambient_prompt"]["unauthorized_risky_action_rate"]),
            pct(summary["ambient_prompt"]["raw_parse_error_rate"]),
        ],
        [
            "Capsule-filtered prompt",
            str(int(summary["capsule_filtered_prompt"]["cases"])),
            pct(summary["capsule_filtered_prompt"]["attack_success_rate"]),
            pct(summary["capsule_filtered_prompt"]["unauthorized_risky_action_rate"]),
            pct(summary["capsule_filtered_prompt"]["raw_parse_error_rate"]),
        ],
        [
            "Absolute reduction",
            str(int(stats["paired_cases"])),
            pct(stats["absolute_attack_success_reduction"]),
            pct(stats["absolute_unauthorized_risky_action_reduction"]),
            f"p ~= {stats['mcnemar_p_value_approx']:.2e}",
        ],
    ]


def capsule_interpretation(key: str) -> str:
    return {
        "intent_capsules": "Balanced result in the tested split.",
        "ablation_no_topic_scope": "Blocks attacks partly by suppressing useful memory behavior.",
        "ablation_no_denied_actions": "No held-out leak here, but stress tests still require this control.",
        "ablation_no_quorum": "No held-out leak here, but quorum is needed for crafted multi-source cases.",
    }[key]


def blocks(results: dict[str, Any], figures: dict[str, Path]) -> list[tuple[str, Any]]:
    return [
        ("abstract", (
            "Long-term-memory LLM agents increasingly store user preferences, tool outputs, web observations, OCR text, "
            "summaries, and prior experiences across sessions. This persistence improves personalization, but it also "
            "creates a durable poisoning channel: a memory inserted during one interaction can later be retrieved as normal "
            "context and steer recommendations, plans, or tool calls. This paper argues that persistent agent poisoning is "
            "not only a malicious-content detection problem; it is an authority-control problem. The proposed defense, "
            "intent-bound memory capsules, compiles each stored memory into a bounded record carrying source, topic, action, "
            "authority, influence-budget, verification, lineage, and freshness constraints. Retrieval is treated as relevance, "
            "not authorization: before a memory can shape planning or action, the capsule contract must permit that influence. "
            "A Python prototype, CapsuleGuard, evaluates this idea against ambient memory, keyword filtering, provenance-only "
            "retrieval, trust-score retrieval, counterfactual memory, and ablated capsule variants. On a held-out workflow-corpus "
            "test split, intent-bound capsules reduced attack success from 33.33% for ambient memory and 24.38% for provenance-only "
            "retrieval to 0.00%, while maintaining 100.00% benign accuracy and 0.00% false positives. Across generated holdout, "
            "attacker-generated, multimodal/OCR-style, trusted-source compromise, and advanced stress suites, the prototype "
            "maintained 0.00% attack success in the tested simulator. A medium live-LLM planner run across llama3, mistral, "
            "and phi3 produced 22.22% ambient-prompt attack success but 0.00% final attack success and 0.00% risky action "
            "under capsule-filtered authorization, with 0.00% raw parse error. These results support least-privilege memory "
            "authorization as a promising defensive layer, while the paper explicitly limits its claim to the tested threat "
            "model and identifies real-tool, real-OCR, and external workflow-trace evaluation as necessary next steps."
        )),
        ("keywords", "LLM agents; memory poisoning; prompt injection; agent security; provenance; least privilege; long-term memory; tool safety"),
        ("h1", "1. Introduction"),
        ("p", (
            "Long-term memory is becoming a core component of LLM agents. Agents use memory to preserve preferences, reuse prior "
            "experience, remember tool results, summarize previous sessions, and maintain continuity across workflows. The same "
            "capability creates a persistent attack surface. If an attacker can cause a harmful or misleading memory to be stored, "
            "that memory can remain dormant, resurface in a later session, and influence an apparently unrelated decision."
        )),
        ("p", (
            "The central problem is that many agent designs let retrieval behave like implicit authorization. Once a memory enters "
            "the prompt context, the planner may treat it as a fact, a preference, an instruction, or action-supporting evidence. "
            "This collapses three different questions into one: Is the memory relevant? Where did it come from? What is it allowed "
            "to influence? Intent-bound memory capsules separate those questions."
        )),
        ("callout", "Core invariant: retrieval is not authorization. A memory may be relevant to the user query and still be unauthorized to change the plan or trigger an action."),
        ("h2", "1.1 Contributions"),
        ("bullets", [
            "A memory-authority formulation of persistent agent poisoning, where the failure is ambient influence over future planning rather than only suspicious text.",
            "A capsule schema that binds memory records to source, topic, action, authority, influence-budget, verification, freshness, and lineage constraints.",
            "A policy gate that separates retrieval, planning support, and action authorization for low-, medium-, and high-risk decisions.",
            "A reproducible Python prototype with baselines, ablations, synthetic workflow corpora, stress suites, trace logging, a trace-corpus importer, a safe tool simulator, a signed local ledger, vector-style retrieval backends, and an LLM planner harness.",
            "Simulator evidence showing 0.00% attack success for the capsule defense across the tested held-out and stress suites while preserving benign utility in the reported runs.",
            "Live LLM planner evidence showing that capsule authorization blocks final poisoned actions across llama3, mistral, and phi3 in the medium workflow-corpus run.",
        ]),
        ("figure", (figures["architecture"], "Figure 1. Intent-bound memory capsule architecture.")),
        ("h1", "2. Problem Statement"),
        ("p", (
            "Given an LLM agent with persistent memory M and a user intent I, the goal is to prevent poisoned or insufficiently "
            "authorized memories from steering future planning or action while preserving useful recall. A poisoning attack succeeds "
            "when an attacker-controlled memory changes the selected recommendation, plan, or tool path toward an attacker-preferred "
            "outcome that would not be chosen under trusted memory alone."
        )),
        ("p", (
            "The research question is: can a memory authorization layer reduce persistent agent poisoning while preserving benign "
            "long-term personalization and recall? The tested hypothesis is that poisoning can be reduced when every memory has an "
            "explicit influence contract and the agent enforces that contract before planning or action."
        )),
        ("h2", "2.1 Formal View"),
        ("p", (
            "Let a memory item m have content c, source s, topic scope T, action restrictions D, authority score a, influence budget b, "
            "verification count v, lineage L, and status q. Let an intent i have requested topic set U, requested action x, and risk class r. "
            "A conventional memory agent often evaluates relevance R(m, i) and then injects the top-k memories into the planner. This paper "
            "adds a second predicate, Auth(m, i, x), that must hold before m can influence the plan. Relevance can be high while authorization "
            "is false."
        )),
        ("p", (
            "The security objective is not to prove that the content is harmless. The objective is to prevent an untrusted or insufficiently "
            "supported memory from gaining more influence than its source, scope, and verification allow. This is why the design is closer to "
            "capability control than to ordinary text classification."
        )),
        ("h2", "2.2 Security Invariants"),
        ("bullets", [
            "I1: A sealed or rejected capsule is never eligible for planning context.",
            "I2: A directive-like memory cannot authorize the action it requests.",
            "I3: A capsule cannot influence an intent outside its topic scope.",
            "I4: A capsule cannot authorize an action in its denied-action set.",
            "I5: Medium- and high-risk actions require source authority above the configured risk floor.",
            "I6: Risky decisions require independent evidence rather than a single retrieved memory.",
            "I7: Derived memories cannot exceed the authority of their parent lineage.",
            "I8: Stale memories lose influence unless fresh verified evidence supports them.",
        ]),
        ("h1", "3. Related Work and Design Gap"),
        ("p", (
            "The local literature set establishes a broad attack surface: prompt sanitization limits, multimodal memory poisoning, "
            "self-reinforcing or persistent agent compromise, poisoning memory or knowledge bases, corrupted web-agent memory, "
            "memory control-flow attacks, query-only memory injection, poisoned experience retrieval, and secure multi-agent memory. "
            "Those papers are valuable because they show that poisoning can enter through many channels and remain effective after "
            "the original interaction has ended."
        )),
        ("p", (
            "The gap addressed here is narrower than the full attack landscape. Many defenses ask whether content looks malicious, "
            "whether a source is known, or whether the final output is unsafe. Intent-bound capsules ask an additional question: "
            "what authority does this memory have over the current task? This turns long-term memory from ambient context into a "
            "controlled security object."
        )),
        ("table", (
            "Table 1. Lessons from the reference literature and how they shape the capsule design.",
            ["Prior attack theme", "Lesson for defense", "Capsule design response"],
            [
                ["Prompt sanitization and red-teaming", "Input-time filtering misses dormant memory attacks.", "Enforce controls at storage, retrieval, planning, and action time."],
                ["Multimodal/OCR poisoning", "Poison can arrive as extracted text rather than normal chat.", "Treat OCR and document-derived memory as low-authority context by default."],
                ["Persistent/zombie agents", "Injected behavior can survive across sessions.", "Use status, sealing, and authorization at every later retrieval."],
                ["Memory/knowledge-base poisoning", "Retrieval can surface attacker-controlled records.", "Separate similarity retrieval from influence authorization."],
                ["Memory control-flow attacks", "Stored content can steer tool chains.", "Gate medium/high-risk actions with denied actions and evidence quorum."],
                ["Experience retrieval attacks", "Past episodes can be weaponized as examples.", "Cap experience influence and require independent support."],
                ["Multi-agent memory", "Derived records need lineage and isolation.", "Propagate parent authority caps and writer/source diversity requirements."],
            ],
        )),
        ("p", (
            "The project therefore does not copy an existing named defense. It extracts a common failure mode from the literature: memory is "
            "too often treated as context rather than as a security-bearing object. The original contribution is the influence contract and the "
            "runtime authorization gate that decides whether a retrieved memory may affect a specific plan or action."
        )),
        ("h1", "4. Threat Model"),
        ("figure", (figures["threat"], "Figure 2. Threat model and trust boundaries.")),
        ("h2", "4.1 Protected Assets"),
        ("bullets", [
            "Recommendation integrity: the agent should not choose an attacker-preferred vendor or option because of unauthorized memory.",
            "User preference integrity: verified user preferences should not be overridden by weak external records.",
            "Action safety: high-impact operations such as email, purchase, transfer, delete, database modification, or private-data sharing require stronger evidence.",
            "Memory lifecycle integrity: derived summaries and experiences should not launder weak authority into trusted context.",
            "Auditability: blocked or sealed memories should remain inspectable instead of silently disappearing.",
        ]),
        ("h2", "4.2 Attacker Capabilities"),
        ("p", (
            "The attacker may influence memories through web content, OCR-extracted text, tool output, agent-generated summaries, "
            "experience logs, query-only interaction, delayed triggers, semantic paraphrases, retrieval collisions, or trusted-looking "
            "but unattested metadata. The attacker tries to cause a future query to retrieve and act on the poisoned memory."
        )),
        ("h2", "4.3 Out-of-Scope Capabilities"),
        ("p", (
            "The prototype does not claim resistance to direct modification of defense code, malicious policy administrators, full "
            "memory database compromise, credential theft, compromised verified identity, arbitrary model-weight compromise, or all "
            "possible adaptive LLM jailbreaks. These are important deployment risks, but they are outside the current benchmark."
        )),
        ("h1", "5. Intent-Bound Memory Capsules"),
        ("p", (
            "A capsule is a memory record plus an influence contract. The contract does not merely describe provenance; it constrains "
            "what the memory may do. A web page can be useful context, but it should not authorize a purchase. A tool result can describe "
            "state, but it should not rewrite user preference. An agent summary can help recall, but it should inherit the authority of "
            "its sources instead of becoming trusted by repetition."
        )),
        ("table", (
            "Table 2. Capsule fields and their security role.",
            ["Capsule field", "Security function"],
            [
                ["source_type", "Classifies origin such as verified record, user declaration, web, tool output, OCR, summary, or experience."],
                ["kind", "Distinguishes fact, preference, observation, experience, and directive-like memory."],
                ["allowed_topics", "Limits influence to matching user intents."],
                ["denied_actions", "Blocks specified actions even when a memory is retrieved."],
                ["source_authority", "Sets the maximum trust a source class can carry."],
                ["influence_budget", "Caps decision weight so one memory cannot dominate by relevance alone."],
                ["verification_count", "Records whether corroborated support exists."],
                ["lineage metadata", "Prevents summaries and derived records from laundering weak sources."],
                ["created_at and freshness notes", "Reduces stale memory influence."],
                ["status", "Keeps active, sealed, and rejected states explicit for audit."],
            ],
        )),
        ("figure", (figures["rules"], "Figure 3. Memory trust and influence rule set.")),
        ("h2", "5.1 Authorization Rules"),
        ("bullets", [
            "Topic scope: a memory may influence only intents with sufficient topic overlap.",
            "Action denial: source classes such as web, OCR, tool output, and agent-derived summaries cannot authorize dangerous actions by default.",
            "Authority floors: medium-risk actions require authority >= 0.55; high-risk actions require authority >= 0.85.",
            "Directive sealing: command-like memories are retained for audit but removed from planning eligibility.",
            "Influence budgets: each memory has a capped contribution to planning support.",
            "Evidence quorum: risky actions require independent support across source and writer diversity, with fresh verified evidence.",
            "Lineage and temporal controls: derived memories inherit authority caps, and stale records lose influence.",
        ]),
        ("h2", "5.2 Algorithms"),
        ("code", (
            "Algorithm 1: CompileMemory(seed)\n"
            "  kind <- classify(seed.content, seed.source_type)\n"
            "  authority <- authority_for(seed.source_type, seed.verified)\n"
            "  authority <- min(authority, parent_authority_cap(seed))\n"
            "  topics <- extract_topic_scope(seed.content)\n"
            "  denied <- source_denied_actions(seed.source_type)\n"
            "  if kind is directive: status <- sealed\n"
            "  if high_risk_claim(seed.content) and authority < 0.85: status <- sealed\n"
            "  budget <- influence_budget(authority, kind, seed.verified)\n"
            "  budget <- temporal_decay(seed.observed_at, budget)\n"
            "  return capsule(content, source, kind, topics, denied, authority, budget, status, lineage)\n\n"
            "Algorithm 2: AuthorizePlan(intent, retrieved_capsules)\n"
            "  eligible <- []\n"
            "  for capsule in retrieved_capsules:\n"
            "      reject if status is not active\n"
            "      reject if topic_overlap(capsule, intent) < threshold\n"
            "      reject if intent.action in capsule.denied_actions\n"
            "      reject if capsule.authority < risk_floor(intent.action_risk)\n"
            "      reject if capsule.kind is directive\n"
            "      eligible.append(capsule)\n"
            "  plan <- planner(intent, eligible)\n"
            "  if plan risk is low: allow\n"
            "  if plan risk is medium/high: require quorum and plan-level authorization\n"
            "  otherwise block or require confirmation"
        )),
        ("h1", "6. Prototype Implementation"),
        ("p", (
            "The prototype, CapsuleGuard, is implemented in Python. It includes capsule compilation, policy gates, baseline agents, "
            "retrieval modes, provenance logging, scenario generation, corpus loading, safe tool traces, and experiment runners. The "
            "default high-volume experiments use a deterministic planner to isolate memory-authorization behavior. The repository "
            "also contains a live LLM planner harness with strict planner-schema output, repair/audit tracking, model-level summaries, "
            "and gap reports. The live LLM results are used as a realism check rather than as the only evidence for the claim."
        )),
        ("table", (
            "Table 3. Implementation components used by the research prototype.",
            ["Component", "Role"],
            [
                ["capsule_guard/models.py", "Capsule schema, memory seed, plans, risks, and traces."],
                ["capsule_guard/compiler.py", "Compiles raw memory into scoped capsules with status, authority, lineage, and decay."],
                ["capsule_guard/rules.py", "Defines trust tiers, risk floors, denied high-risk actions, and capsule eligibility."],
                ["capsule_guard/policy.py", "Filters eligible capsules and enforces evidence quorum and plan authorization."],
                ["capsule_guard/provenance.py", "HMAC-backed append-only local provenance ledger for tamper-evidence experiments."],
                ["capsule_guard/vector_backend.py", "SQLite-backed hashed-vector style retrieval backend for retrieval-mode testing."],
                ["capsule_guard/tools.py", "Safe tool simulator and full action trace logging."],
                ["capsule_guard/adaptive_attacker.py", "Bounded closed-loop attacker that mutates attacks after policy feedback."],
                ["run_capsuleguard.py", "Main benchmark runner producing CSV, JSONL traces, tool traces, and charts."],
            ],
        )),
        ("h2", "6.1 Baselines"),
        ("table", (
            "Table 4. Compared agents and what each one tests.",
            ["Agent", "Defense idea", "Expected weakness"],
            [
                ["ambient_memory", "Retrieve memory and let it directly influence planning.", "No security boundary between recall and action."],
                ["keyword_filter", "Seal obvious directive-like phrases.", "Benign-looking poison and paraphrases survive."],
                ["provenance_only", "Score memory by source trust.", "Known or plausible sources can still lack authority for the current action."],
                ["trust_score_retrieval", "Weight retrieval by trust score.", "Trust-weighted relevance is still not plan authorization."],
                ["counterfactual_memory", "Compare planning with and without top memories.", "Can block useful legitimate preference changes, increasing false positives."],
                ["intent_capsules", "Authorize every memory before planning/action.", "Main proposed defense."],
            ],
        )),
        ("h2", "6.2 Attack Coverage"),
        ("table", (
            "Table 5. Attack areas represented in the current simulator.",
            ["Attack area", "Current representation", "Remaining realism gap"],
            [
                ["Adaptive attackers", "Closed-loop mutation after policy feedback.", "Not a fully autonomous external LLM red team."],
                ["Delayed trigger poisoning", "Memories that activate under later task context.", "Trigger text is still synthetic."],
                ["Cross-session poisoning", "Workflow and trace-corpus records that persist into later planning.", "No multi-day production trace set yet."],
                ["Tool-chain manipulation", "Safe tool simulator and action traces.", "No real browser, email, or database side effects."],
                ["Semantic paraphrase poisoning", "Synonym and alias variants plus generated workflow variants.", "No large external paraphrase corpus yet."],
                ["Retrieval collision attacks", "Similarity-overlap stress cases.", "No industrial embedding collision search yet."],
                ["Multimodal hidden instruction poisoning", "OCR/document-derived text cases.", "No raw-image hidden-pixel pipeline yet."],
                ["Trusted-source compromise", "Unattested and trusted-looking metadata cases.", "No real identity-provider compromise model."],
            ],
        )),
        ("h1", "7. Evaluation Methodology"),
        ("h2", "7.1 Research Questions"),
        ("bullets", [
            "RQ1: Does memory authorization reduce attack success compared with ambient memory, keyword filtering, and provenance-only baselines?",
            "RQ2: Does the defense preserve benign memory usefulness?",
            "RQ3: Which capsule controls matter under ablation?",
            "RQ4: Which attack classes remain difficult for provenance-only and counterfactual baselines?",
            "RQ5: When a live LLM planner is used, does capsule filtering prevent tempted poisoned plans from surviving final authorization?",
        ]),
        ("h2", "7.2 Benchmarks and Metrics"),
        ("p", (
            "The primary benchmark is a held-out workflow-corpus test split with disjoint train, development, and test records. The "
            "split contains 60 train, 24 development, and 36 held-out test records across vendor recommendation, email, calendar, "
            "file search, CRM notes, web research, OCR-style documents, and tool-chain workflows. Supporting suites evaluate generated "
            "holdout attacks, attacker-generated variants, multimodal/OCR-style attacks, trusted-source compromise, and advanced stress cases."
        )),
        ("bullets", [
            "Attack Success Rate (ASR): fraction of poisoned cases that produce the attacker target.",
            "Unauthorized risky action rate: fraction of poisoned cases producing unauthorized medium/high-risk behavior.",
            "Benign accuracy: fraction of benign tasks answered correctly.",
            "False positive rate (FPR): fraction of benign cases incorrectly blocked.",
            "Poison sealing rate: fraction of poison cases sealed at compile time rather than merely blocked later.",
            "Planner attack success: live LLM chose the attacker target before final authorization.",
            "Raw parse error rate: first LLM planner output failed the strict schema before any repair.",
        ]),
        ("p", (
            "The evaluation reports both security and utility metrics because a trivial defense can achieve low ASR by ignoring memory entirely. "
            "The useful target is the intersection of low ASR, low unauthorized action rate, high benign accuracy, and low FPR. This is why the "
            "ablation without topic scope is treated as weaker even though it reports 0.00% ASR on the held-out split."
        )),
        ("h1", "8. Results"),
        ("figure", (figures["workflow"], "Figure 4. Held-out workflow-corpus results.")),
        ("table", (
            "Table 6. Held-out workflow-corpus test split.",
            ["Agent", "ASR", "Risky action", "Benign accuracy", "FPR", "Sealing"],
            workflow_table(results),
        )),
        ("p", (
            "On the held-out workflow corpus, ambient memory and keyword filtering both reached 33.33% ASR. Provenance-only and "
            "trust-score retrieval reduced ASR to 24.38% but still allowed attacks through. Counterfactual memory reduced ASR to "
            "3.96%, but benign accuracy dropped to 91.25% and FPR rose to 8.75%. The full capsule defense reached 0.00% ASR, "
            "0.00% unauthorized risky action rate, 100.00% benign accuracy, and 0.00% FPR."
        )),
        ("figure", (figures["stress"], "Figure 5. Stress-suite attack success rates.")),
        ("table", (
            "Table 7. Stress-suite ASR comparison.",
            ["Scenario", "Ambient ASR", "Provenance ASR", "Capsule ASR", "Capsule sealing"],
            stress_table(results),
        )),
        ("p", (
            "The stress suites show the same pattern: provenance helps but does not fully solve poisoning. Trusted-source compromise "
            "is especially important: provenance-only ASR remains 28.11%, while the capsule defense reports 0.00% ASR in the tested "
            "simulator. This supports the paper's main point that source labels alone are not enough; memories must be authorized for "
            "the particular influence they attempt to exert."
        )),
        ("figure", (figures["ablation"], "Figure 6. Ablation results.")),
        ("table", (
            "Table 8. Ablation results on the held-out workflow corpus.",
            ["Variant", "ASR", "Benign", "FPR", "Sealing", "Interpretation"],
            ablation_table(results),
        )),
        ("p", (
            "The ablation results show why the paper should not report ASR alone. Removing topic scope still yields 0.00% ASR in this "
            "held-out split, but benign accuracy falls to 53.75%. That is not a good defense; it blocks attacks partly by suppressing "
            "useful memory behavior. The full capsule configuration is stronger because it combines zero tested attack success with "
            "full benign accuracy and zero false positives."
        )),
        ("h2", "8.1 Live LLM Planner Check"),
        ("figure", (figures["live_llm"], "Figure 7. Live LLM planner check.")),
        ("table", (
            "Table 9. Medium live LLM workflow-corpus run.",
            ["Condition", "Rows", "Planner tempted", "Final ASR", "Risky action", "Raw parse error", "Final parse error"],
            live_llm_table(results),
        )),
        ("p", (
            "The live LLM planner experiment is not the highest-volume benchmark; it is a realism check that the authorization layer works "
            "when the plan is produced by actual local models rather than only by the deterministic planner. Across llama3, mistral, and "
            "phi3, the ambient prompt produced 22.22% final attack success. The capsule-filtered prompt still showed 2.78% planner temptation, "
            "which is useful evidence that the planner can be pulled toward the poison, but final authorization reduced accepted attack success "
            "and risky action to 0.00%. Raw and final parse errors were both 0.00%, so this result is not explained by malformed model output."
        )),
        ("table", (
            "Table 10. Defended medium live LLM result by model.",
            ["Model", "Rows", "Planner tempted", "Final ASR", "Risky action", "Raw parse error"],
            live_llm_model_table(results),
        )),
        ("table", (
            "Table 11. High-cost local LLM smoke profile.",
            ["Condition", "Rows", "Final ASR", "Risky action", "Parse/audit note"],
            high_cost_smoke_table(results),
        )),
        ("p", (
            "The high-cost local smoke profile is a machinery check for the larger conference-grade evaluation path. It uses local simulated "
            "planner providers to exercise the high-cost case sampler, paired statistics, raw-output audit logs, and gap reports. It should not "
            "be reported as paid API evidence, but it shows that the larger evaluation harness is ready to run against additional live or paid models."
        )),
        ("h2", "8.2 Gap-Closure Interpretation"),
        ("table", (
            "Table 12. Why capsule authorization closes tested baseline failures.",
            ["Failure class", "Why simpler baselines fail", "Capsule control that blocks it"],
            [
                ["Benign-looking web poison", "Keyword filters see ordinary recommendation language.", "Source authority floor plus denied high-risk actions."],
                ["Experience poisoning", "Provenance treats experience logs as plausible prior context.", "Influence cap plus independent evidence quorum."],
                ["Agent-summary laundering", "Derived summaries can hide weak source origins.", "Lineage authority cap and agent-derived restrictions."],
                ["Recency pressure", "Recent poisoned memory can outrank older verified preference.", "Verified preference precedence plus freshness-aware influence."],
                ["Semantic paraphrase", "Attack avoids exact suspicious words.", "Authorization checks action/topic/authority instead of only keywords."],
                ["Trusted-looking metadata", "Source labels alone can be spoofed in the simulator.", "Attestation downgrade and independent writer/source requirements."],
                ["Tool-output steering", "Tool results can look operationally relevant.", "Tool outputs cannot authorize dangerous actions without stronger support."],
            ],
        )),
        ("h2", "8.3 Security Argument"),
        ("p", (
            "The capsule defense succeeds in the tested cases because the attacker must satisfy multiple independent conditions at once. A poison "
            "memory must be relevant enough to retrieve, active rather than sealed, within topic scope, outside the denied-action set, above the "
            "risk-specific authority floor, fresh or sufficiently supported, and part of an independent quorum for risky action. A failure in any "
            "one of these checks prevents that memory from becoming action-authorizing evidence."
        )),
        ("p", (
            "This is not a mathematical proof of universal security. It is an engineering security argument supported by ablation and stress-test "
            "evidence. The result is strongest where attackers can write or influence memory but cannot modify policy, forge all independent trusted "
            "writers, or compromise the verification source."
        )),
        ("h1", "9. Discussion"),
        ("p", (
            "The experiments support three design lessons. First, suspicious-text filtering is too shallow because poisoned memory can "
            "look like ordinary preference, experience, or tool context. Second, provenance improves robustness but does not answer what "
            "a memory is authorized to influence. Third, action gates need memory-level evidence, not just final output moderation, because "
            "the unsafe influence may already have shaped the plan before the final response is checked."
        )),
        ("p", (
            "The relatively low poison sealing rate should not be misread as failure. Sealing is an early quarantine mechanism for obvious "
            "directive or high-risk weak-source records. Many subtler poison records are intentionally left inspectable but blocked later by "
            "topic scope, authority floors, influence caps, lineage checks, and evidence quorum. The architecture therefore uses both compile-time "
            "quarantine and runtime authorization."
        )),
        ("p", (
            "A practical deployment would use the capsule layer alongside conventional defenses: prompt sanitization at intake, signed provenance "
            "at write time, access-controlled retrieval, output validation, human confirmation for irreversible actions, and tool sandboxing. The "
            "capsule layer is valuable because it protects the middle of the memory lifecycle, where poisoned data is already stored but has not yet "
            "been allowed to influence a plan."
        )),
        ("h1", "10. Limitations and Threats to Validity"),
        ("bullets", [
            "The main evaluation is simulator-based, not a deployed agent study.",
            "The default high-volume planner is deterministic; the live LLM planner suite is now reported, but it remains medium-scale rather than a large multi-seed model study.",
            "The workflow corpus is generated and curated. The repo now includes a trace-corpus importer, but external redacted agent traces still need to be collected and reported.",
            "The vector backend is a hashed/SQLite approximation, not a full industrial embedding retrieval system with adversarial collision search.",
            "OCR and multimodal attacks are represented mainly as extracted text rather than raw image, hidden-pixel, or alt-text pipelines.",
            "Source labels are modeled as metadata; production systems need signed identities, append-only provenance, and administrative controls.",
            "Policy thresholds are hand-tuned research parameters and require workload calibration.",
            "The prototype does not claim security against database compromise, malicious policy administrators, stolen credentials, or arbitrary model compromise.",
        ]),
        ("h2", "10.1 Industrialization Roadmap"),
        ("bullets", [
            "Replace generated workflow records with redacted real or lab-user traces.",
            "Scale the LLM planner harness beyond the current medium run: more models, more seeds, more task domains, and paid API models if available.",
            "Use a production vector database and explicit adversarial embedding-collision search.",
            "Bind source identity to cryptographic signatures or enterprise identity providers.",
            "Run raw-image OCR, alt-text, and document-ingestion tests instead of only extracted text.",
            "Calibrate thresholds on benign production workloads before using them as deployment policy.",
            "Add isolated browser, email, database, and payment sandboxes for realistic tool-side-effect testing.",
        ]),
        ("h1", "11. Ethics and Safety"),
        ("p", (
            "This work is defensive. The benchmark uses synthetic vendors, synthetic memories, and safe tool simulations. It does not include "
            "real credentials, real private data, or instructions for compromising deployed systems. Attack examples are limited to the level "
            "needed to evaluate the defense in a controlled setting."
        )),
        ("h1", "12. Conclusion"),
        ("p", (
            "Persistent agent poisoning should be treated as a memory lifecycle and authority-control problem. Intent-bound memory capsules "
            "make the authority of each memory explicit and enforce it before planning or action. In the tested simulator, the capsule defense "
            "reduced attack success to 0.00% across the reported suites while preserving benign utility in the primary held-out workflow test. "
            "In the medium live-LLM workflow-corpus run, capsule-filtered authorization also kept final ASR and risky action at 0.00% across "
            "llama3, mistral, and phi3 while preserving parse validity. The supported claim is deliberately bounded: least-privilege memory "
            "authorization is a promising layer for defending LLM agents with long-term memory. The next step is larger live-model evaluation, "
            "real vector databases, real OCR pipelines, sandboxed tools, and external workflow traces."
        )),
        ("h1", "References"),
        ("references", references()),
        ("h1", "Appendix A. Reproducibility Commands"),
        ("code", (
            "Repository snapshot: C:\\Users\\User\\Music\\Agent-Poisoning-Research-FINAL\n"
            "Branch: fix/llm-first-pass-json-planning\n"
            "Commit: ac1d2ea\n\n"
            "python -m unittest discover -s tests\n"
            "# Expected current result: Ran 136 tests, OK\n\n"
            "python run_capsuleguard.py --attack-mode workflow_corpus "
            "--workflow-corpus data/workflow_corpus_splits/test.jsonl "
            "--trials 5 --repetitions 4 --noise-memories 4 --seed 2026 "
            "--summary-csv results/workflow_corpus_test_split_summary.csv "
            "--trace-jsonl results/workflow_corpus_test_split_traces.jsonl "
            "--breakdown-csv results/workflow_corpus_test_split_breakdown.csv "
            "--gap-closure-csv results/workflow_corpus_test_split_gap_closure.csv "
            "--tool-trace-csv results/workflow_corpus_test_split_tool_traces.csv "
            "--charts-dir results/workflow_corpus_test_split_charts\n\n"
            "python run_capsuleguard.py --attack-mode trace_corpus "
            "--workflow-corpus data/agent_trace_corpus_sample.jsonl "
            "--summary-csv results/trace_corpus_sample_summary.csv "
            "--trace-jsonl results/trace_corpus_sample_traces.jsonl "
            "--breakdown-csv results/trace_corpus_sample_breakdown.csv "
            "--gap-closure-csv results/trace_corpus_sample_gap_closure.csv "
            "--tool-trace-csv results/trace_corpus_sample_tool_traces.csv "
            "--charts-dir results/trace_corpus_sample_charts\n\n"
            "python run_llm_experiment.py --provider ollama "
            "--models llama3,mistral,phi3 "
            "--case-source workflow-corpus "
            "--workflow-corpus data/workflow_corpus_splits/test.jsonl "
            "--case-limit 36 --case-seed 2026 --repetitions 1 "
            "--output-csv results/gap_fix_medium_live_llm_suite.csv "
            "--summary-csv results/gap_fix_medium_live_llm_summary.csv "
            "--model-summary-csv results/gap_fix_medium_live_llm_model_summary.csv "
            "--gap-report-csv results/gap_fix_medium_live_llm_gap_report.csv"
        )),
        ("callout", (
            "Submission note: before formal submission, complete bibliographic metadata for each reference and rerun the final experiments "
            "from the public repository snapshot. Do not claim that the method solves all agent memory poisoning."
        )),
    ]


def references() -> list[str]:
    return [
        "Mohamed Amine Ferrag, Abderrahmane Lakas, Norbert Tihanyi, and Merouane Debbah. 2025. Securing LLM agents: From prompt sanitization to autonomous red teaming and beyond. Internet of Things and Cyber-Physical Systems 5, 185-209. doi:10.1016/j.iotcps.2026.03.001.",
        "Jiachen Qian. 2026. Visual Inception: Compromising Long-term Planning in Agentic Recommenders via Multimodal Memory Poisoning. arXiv:2604.16966.",
        "Xianglin Yang, Yufei He, Shuo Ji, Bryan Hooi, and Jin Song Dong. 2026. Zombie Agents: Persistent Control of Self-Evolving LLM Agents via Self-Reinforcing Injections. arXiv:2602.15654.",
        "Zhaorun Chen, Zhen Xiang, Chaowei Xiao, Dawn Song, and Bo Li. 2024. AgentPoison: Red-teaming LLM Agents via Poisoning Memory or Knowledge Bases. arXiv:2407.12784.",
        "Atharv Singh Patlan, Ashwin Hebbar, Pramod Viswanath, and Prateek Mittal. 2025. Context manipulation attacks: Web agents are susceptible to corrupted memory. arXiv:2506.17318.",
        "Zhenlin Xu, Xiaogang Zhu, Yu Yao, Minhui Xue, and Yiliao Song. 2026. From Storage to Steering: Memory Control Flow Attacks on LLM Agents. arXiv:2603.15125.",
        "Hanling Tian, Zeyang Sha, Jingying Wang, Yuhang Liu, Zhehao Huang, and Xiaolin Huang. 2025. InjecMEM: Memory Injection Attack on LLM Agent Memory Systems. OpenReview submission to ICLR 2026.",
        "Vicenc Torra and Maria Bras-Amoros. 2026. Memory poisoning and secure multi-agent systems. arXiv:2603.20357.",
        "Balachandra Devarangadi Sunil, Isheeta Sinha, Piyush Maheshwari, Shantanu Todmal, Shreyan Mallik, and Shuchi Mishra. 2026. Memory Poisoning Attack and Defense on Memory Based LLM-Agents. arXiv:2601.05504.",
        "Saksham Sahai Srivastava and Haoyu He. 2025. MemoryGraft: Persistent Compromise of LLM Agents via Poisoned Experience Retrieval. arXiv:2512.16962.",
        "Shen Dong, Shaochen Xu, Pengfei He, Yige Li, Jiliang Tang, Tianming Liu, Hui Liu, and Zhen Xiang. 2025. Memory Injection Attacks on LLM Agents via Query-Only Interaction. arXiv:2503.03704.",
        "OWASP Foundation. OWASP Top 10 for Large Language Model Applications, 2025.",
    ]


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |"]
    out.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        out.append("| " + " | ".join(str(cell).replace("\n", " ") for cell in row) + " |")
    return "\n".join(out)


def write_markdown(results: dict[str, Any], figures: dict[str, Path]) -> None:
    parts = [
        f"# {TITLE}",
        "",
        AUTHOR,
        AFFILIATION,
        DATE_LINE,
        "",
    ]
    for kind, payload in blocks(results, figures):
        if kind == "abstract":
            parts.extend(["## Abstract", "", payload, ""])
        elif kind == "keywords":
            parts.extend(["**Keywords:** " + payload, ""])
        elif kind == "h1":
            parts.extend([f"## {payload}", ""])
        elif kind == "h2":
            parts.extend([f"### {payload}", ""])
        elif kind == "p":
            parts.extend([payload, ""])
        elif kind == "callout":
            parts.extend(["> " + payload, ""])
        elif kind == "bullets":
            parts.extend([*(f"- {item}" for item in payload), ""])
        elif kind == "code":
            parts.extend(["```text", payload, "```", ""])
        elif kind == "figure":
            path, caption = payload
            parts.extend([f"![{caption}]({path.relative_to(ROOT).as_posix()})", "", f"*{caption}*", ""])
        elif kind == "table":
            caption, headers, rows = payload
            parts.extend([f"**{caption}**", "", markdown_table(headers, rows), ""])
        elif kind == "references":
            for idx, ref in enumerate(payload, 1):
                parts.append(f"[{idx}] {ref}")
            parts.append("")
    OUT_MD.write_text("\n".join(parts), encoding="utf-8")


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*hex_rgb(color))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color: str = GRID, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    borders.append(bottom)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.lstrip("#"))


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_fixed(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)
    return p


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "#F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_code_block(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "#F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(text.splitlines()):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(5)
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(caption)
    set_run_font(run, size=9.5, color=BLUE, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_fixed(table)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_shading(hdr[idx], LIGHT_GRAY)
        set_cell_margins(hdr[idx])
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=8.7, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=8.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=9, color=GRAY, italic=True)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    p.paragraph_format.space_before = Pt(13 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 11.5, color=BLUE if level == 1 else NAVY, bold=True)


def build_docx(results: dict[str, Any], figures: dict[str, Path]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Intent-Bound Memory Capsules")
    set_run_font(run, size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Draft for research development - May 2026")
    set_run_font(run, size=8.5, color=GRAY)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(5)
    title_run = title_p.add_run(TITLE)
    set_run_font(title_run, size=18, color=NAVY, bold=True)

    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(1)
    ar = author_p.add_run(f"{AUTHOR}\n{AFFILIATION}\n{DATE_LINE}")
    set_run_font(ar, size=10.5, color=BLACK)
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, GRID, "8")

    for kind, payload in blocks(results, figures):
        if kind == "abstract":
            add_heading(doc, "Abstract", 1)
            add_body_paragraph(doc, payload)
        elif kind == "keywords":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            r1 = p.add_run("Keywords: ")
            set_run_font(r1, size=10.5, color=BLACK, bold=True)
            r2 = p.add_run(payload)
            set_run_font(r2, size=10.5, color=BLACK)
        elif kind == "h1":
            add_heading(doc, payload, 1)
        elif kind == "h2":
            add_heading(doc, payload, 2)
        elif kind == "p":
            add_body_paragraph(doc, payload)
        elif kind == "callout":
            add_callout(doc, payload)
        elif kind == "bullets":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(item)
                set_run_font(run, size=10.2, color=BLACK)
        elif kind == "code":
            add_code_block(doc, payload)
        elif kind == "figure":
            path, caption = payload
            add_figure(doc, path, caption)
        elif kind == "table":
            caption, headers, rows = payload
            add_table(doc, caption, headers, rows)
        elif kind == "references":
            for idx, ref in enumerate(payload, 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(f"[{idx}] {ref}")
                set_run_font(run, size=9.5, color=BLACK)

    doc.save(OUT_DOCX)


def main() -> None:
    ROOT.mkdir(parents=True, exist_ok=True)
    if FIG_DIR.exists():
        for old_png in FIG_DIR.glob("*.png"):
            old_png.unlink()
    results = load_all_results()
    write_tables_csv(results)
    figures = build_figures(results)
    write_markdown(results, figures)
    build_docx(results, figures)
    print(f"wrote {OUT_MD}")
    print(f"wrote {OUT_DOCX}")
    print(f"wrote {OUT_TABLES}")
    print(f"wrote {FIG_DIR}")


if __name__ == "__main__":
    main()
